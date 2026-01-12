"""
Script para agregar ÚNICAMENTE el Epígrafe VI al documento existente
Ingeniería de requisitos del sistema de gestión de pedidos
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def epigrafe_vi(doc):
    """Genera el Epígrafe VI: Ingeniería de requisitos del sistema de gestión de pedidos"""
    # EPÍGRAFE VI
    p = doc.add_paragraph('VI. Ingeniería de requisitos del sistema de gestión de pedidos')
    try:
        p.style = 'Subtitulo'
    except:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)
    
    intro_requisitos = """El presente epígrafe documenta los requisitos funcionales y no funcionales del sistema de gestión de pedidos, así como las historias de usuario que guiaron el desarrollo iterativo mediante Extreme Programming (XP). Los requisitos fueron identificados mediante las técnicas de recopilación descritas en el Epígrafe III y priorizados según su impacto en el negocio y complejidad técnica. La clasificación de requisitos no funcionales sigue la taxonomía propuesta por Sommerville (2011), garantizando la cobertura integral de atributos de calidad del software."""
    
    doc.add_paragraph(intro_requisitos)
    
    # 6.1 Requisitos funcionales
    doc.add_paragraph()
    p = doc.add_paragraph('6.1 Requisitos funcionales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_6_1 = """Los requisitos funcionales especifican las capacidades y comportamientos que el sistema debe exhibir para satisfacer las necesidades del negocio. Se identificaron 44 requisitos funcionales, organizados en ocho módulos principales: autenticación y gestión de usuarios (8 RF), gestión de clientes (5 RF), gestión de productos e inventario (6 RF), gestión de pedidos (7 RF), gestión de pagos (6 RF), reportes y estadísticas (5 RF), auditoría (2 RF), y gestión de devoluciones (4 RF). La Tabla 6.1 presenta un resumen de los requisitos funcionales más relevantes."""
    
    doc.add_paragraph(texto_6_1)
    doc.add_paragraph()
    
    # Tabla de RF (muestra resumida de los más importantes)
    p = doc.add_paragraph('Tabla 6.1: Requisitos funcionales principales del sistema')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=13, cols=5)
    table.style = 'Light Grid Accent 1'
    
    headers = ['No.', 'Nombre', 'Descripción', 'Prioridad', 'Complejidad']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    requisitos = [
        ['RF-01', 'Crear usuario interno', 'Permitir registrar usuarios con rol (admin, supervisor, vendedor)', 'Alta', 'Media'],
        ['RF-07', 'Iniciar sesión', 'Autenticación segura con JWT y bcrypt', 'Alta', 'Media'],
        ['RF-09', 'Crear cliente', 'Registrar clientes con datos y contactos', 'Alta', 'Media'],
        ['RF-14', 'Crear producto', 'Registrar productos con precio y stock', 'Alta', 'Media'],
        ['RF-17', 'Listar productos (catálogo público)', 'Mostrar productos disponibles sin autenticación', 'Alta', 'Baja'],
        ['RF-19', 'Reducir stock automáticamente', 'Actualizar inventario al confirmar pedido', 'Alta', 'Alta'],
        ['RF-21', 'Crear pedido', 'Registrar pedido con múltiples productos', 'Alta', 'Alta'],
        ['RF-22', 'Validar stock disponible', 'Verificar existencias antes de confirmar', 'Alta', 'Media'],
        ['RF-28', 'Registrar pago', 'Permitir pagos parciales acumulativos', 'Alta', 'Alta'],
        ['RF-30', 'Actualizar total pagado', 'Sumar pagos y cambiar estado automáticamente', 'Alta', 'Alta'],
        ['RF-34', 'Estadísticas diarias', 'Reporte de ventas del día por método de pago', 'Media', 'Media'],
        ['RF-41', 'Registrar devolución', 'Cambiar estado y eliminar pagos asociados', 'Media', 'Alta']
    ]
    
    for i, req in enumerate(requisitos, start=1):
        row_cells = table.rows[i].cells
        for j, dato in enumerate(req):
            row_cells[j].text = dato
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    texto_rf_nota = """La lista completa de 44 requisitos funcionales se encuentra documentada en el sistema de gestión de requisitos del proyecto, incluyendo criterios de aceptación detallados para cada uno."""
    p = doc.add_paragraph(texto_rf_nota)
    p.runs[0].italic = True
    
    # 6.2 Requisitos no funcionales
    doc.add_paragraph()
    p = doc.add_paragraph('6.2 Requisitos no funcionales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_6_2 = """Los requisitos no funcionales especifican atributos de calidad, restricciones y propiedades del sistema que no están directamente relacionadas con funcionalidades específicas, pero que determinan su viabilidad, usabilidad y aceptación. Se identificaron 33 requisitos no funcionales, clasificados según la taxonomía de Sommerville (2011) en ocho categorías: rendimiento (4 RNF), seguridad (10 RNF), usabilidad (4 RNF), compatibilidad (4 RNF), mantenibilidad (3 RNF), fiabilidad y disponibilidad (4 RNF), escalabilidad (3 RNF), y legalidad y privacidad (3 RNF)."""
    
    doc.add_paragraph(texto_6_2)
    doc.add_paragraph()
    
    # Subsecciones de RNF
    categorias_rnf = [
        ('6.2.1 Requisitos de rendimiento', [
            'RNF-01: El catálogo de productos debe cargarse en menos de 3 segundos.',
            'RNF-02: Las búsquedas deben ejecutarse en menos de 2 segundos.',
            'RNF-03: El sistema debe soportar al menos 500 usuarios simultáneos.',
            'RNF-04: Los reportes deben generarse en menos de 10 segundos.'
        ]),
        ('6.2.2 Requisitos de seguridad', [
            'RNF-05: Las contraseñas deben almacenarse cifradas con bcrypt.',
            'RNF-06: La comunicación debe usar HTTPS en producción.',
            'RNF-07: Solo usuarios autenticados pueden acceder al sistema (excepto catálogo).',
            'RNF-08: El sistema debe registrar todas las acciones en logs de auditoría.',
            'RNF-09: Los roles deben validarse en backend para evitar manipulación.',
            'RNF-10: Debe permitirse recuperación de contraseña mediante correo electrónico.'
        ]),
        ('6.2.3 Requisitos de usabilidad', [
            'RNF-11: La interfaz debe ser intuitiva para usuarios no técnicos.',
            'RNF-12: El sistema debe ser responsive (móviles, tabletas, PC).',
            'RNF-13: Los botones deben estar etiquetados claramente.',
            'RNF-14: Deben mostrarse confirmaciones visuales de las acciones.'
        ]),
        ('6.2.4 Requisitos de compatibilidad', [
            'RNF-16: La API debe ser compatible con cualquier cliente HTTP.',
            'RNF-17: Debe integrarse con PostgreSQL 14 o superior.',
            'RNF-18: Debe permitir integración con servicios de notificación.',
            'RNF-19: Debe exportar información en formatos PDF, XLSX, CSV.'
        ]),
        ('6.2.5 Requisitos de mantenibilidad', [
            'RNF-20: El código debe estar documentado y versionado en Git.',
            'RNF-21: El código debe seguir arquitectura modular en capas.',
            'RNF-22: Debe permitir agregar nuevos roles sin modificar estructura base.',
            'RNF-23: Las actualizaciones deben realizarse sin pérdida de datos.'
        ]),
        ('6.2.6 Requisitos de fiabilidad y disponibilidad', [
            'RNF-24: El sistema debe garantizar disponibilidad del 99% mensual.',
            'RNF-25: Debe tener respaldos automáticos diarios.',
            'RNF-26: Debe recuperarse en menos de 10 minutos ante fallas.',
            'RNF-27: Debe evitar pérdida de datos ante desconexiones inesperadas.'
        ]),
        ('6.2.7 Requisitos de escalabilidad', [
            'RNF-28: Debe permitir ampliación para más productos, usuarios y tiendas.',
            'RNF-29: Debe permitir conexión con futuras apps móviles.',
            'RNF-30: La arquitectura debe soportar crecimiento sin reestructuración.'
        ]),
        ('6.2.8 Requisitos de legalidad y privacidad', [
            'RNF-31: Debe cumplir con leyes de protección de datos (GDPR o equivalente).',
            'RNF-32: Los datos de clientes no pueden compartirse sin consentimiento.',
            'RNF-33: Debe permitir eliminar cuentas y datos personales si se solicita.'
        ])
    ]
    
    for titulo, requisitos_lista in categorias_rnf:
        doc.add_paragraph()
        p = doc.add_paragraph(titulo)
        p.runs[0].bold = True
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(12)
        
        for req in requisitos_lista:
            doc.add_paragraph(req, style='List Bullet')
    
    # 6.3 Historias de usuario
    doc.add_paragraph()
    p = doc.add_paragraph('6.3 Historias de usuario')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_6_3 = """Las historias de usuario representan funcionalidades desde la perspectiva de los usuarios finales, expresadas en lenguaje natural siguiendo el formato: "Como [rol] quiero [funcionalidad] para [objetivo]". Se definieron 11 historias de usuario principales, priorizadas según valor de negocio y agrupando múltiples requisitos funcionales relacionados. Cada historia incluye criterios de aceptación, tiempo estimado de implementación (en puntos de historia) y prioridad."""
    
    doc.add_paragraph(texto_6_3)
    doc.add_paragraph()
    
    # Tabla de HU
    p = doc.add_paragraph('Tabla 6.2: Historias de usuario del sistema')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_hu = doc.add_table(rows=12, cols=4)
    table_hu.style = 'Light Grid Accent 1'
    
    headers_hu = ['ID', 'Historia de Usuario', 'Prioridad', 'Puntos']
    header_cells_hu = table_hu.rows[0].cells
    for i, header in enumerate(headers_hu):
        header_cells_hu[i].text = header
        for paragraph in header_cells_hu[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    historias = [
        ['HU-01', 'Como administrador quiero gestionar usuarios y roles para controlar el acceso al sistema', 'Alta', '8'],
        ['HU-02', 'Como usuario quiero autenticarme de forma segura para acceder a mis funciones', 'Alta', '5'],
        ['HU-03', 'Como vendedor quiero gestionar clientes para mantener actualizada su información', 'Alta', '8'],
        ['HU-04', 'Como administrador quiero gestionar productos e inventario para controlar existencias', 'Alta', '13'],
        ['HU-05', 'Como cliente quiero ver el catálogo público de productos para conocer la oferta', 'Media', '3'],
        ['HU-06', 'Como vendedor quiero crear y gestionar pedidos para registrar las ventas', 'Alta', '13'],
        ['HU-07', 'Como vendedor quiero registrar pagos parciales o totales para llevar control financiero', 'Alta', '8'],
        ['HU-08', 'Como administrador quiero consultar reportes y estadísticas para analizar el desempeño', 'Media', '8'],
        ['HU-09', 'Como auditor quiero consultar logs de acciones para rastrear operaciones', 'Media', '5'],
        ['HU-10', 'Como vendedor quiero registrar devoluciones para revertir pedidos cuando sea necesario', 'Media', '8'],
        ['HU-11', 'Como administrador quiero buscar clientes por múltiples criterios para localizar información rápidamente', 'Baja', '5']
    ]
    
    for i, hu in enumerate(historias, start=1):
        row_cells = table_hu.rows[i].cells
        for j, dato in enumerate(hu):
            row_cells[j].text = dato
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # 6.4 Agrupación de RF por HU
    doc.add_paragraph()
    p = doc.add_paragraph('6.4 Trazabilidad entre historias de usuario y requisitos funcionales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_6_4 = """La trazabilidad entre historias de usuario y requisitos funcionales garantiza que cada funcionalidad implementada responde a una necesidad del usuario y que todos los requisitos están cubiertos por al menos una historia. La Tabla 6.3 presenta la matriz de trazabilidad."""
    
    doc.add_paragraph(texto_6_4)
    doc.add_paragraph()
    
    # Tabla de trazabilidad
    p = doc.add_paragraph('Tabla 6.3: Matriz de trazabilidad HU ↔ RF')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table_traz = doc.add_table(rows=12, cols=2)
    table_traz.style = 'Light Grid Accent 1'
    
    headers_traz = ['Historia de Usuario', 'Requisitos Funcionales Asociados']
    header_cells_traz = table_traz.rows[0].cells
    for i, header in enumerate(headers_traz):
        header_cells_traz[i].text = header
        for paragraph in header_cells_traz[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    trazabilidad = [
        ['HU-01', 'RF-01, RF-02, RF-03, RF-04, RF-05, RF-06'],
        ['HU-02', 'RF-07, RF-08'],
        ['HU-03', 'RF-09, RF-10, RF-11, RF-12, RF-13, RF-44'],
        ['HU-04', 'RF-14, RF-15, RF-16, RF-17, RF-18, RF-19, RF-20'],
        ['HU-05', 'RF-18 (catálogo público)'],
        ['HU-06', 'RF-21, RF-22, RF-23, RF-24, RF-25, RF-26, RF-27'],
        ['HU-07', 'RF-28, RF-29, RF-30, RF-31, RF-32, RF-33'],
        ['HU-08', 'RF-34, RF-35, RF-36, RF-37, RF-38'],
        ['HU-09', 'RF-39, RF-40'],
        ['HU-10', 'RF-41, RF-42, RF-43'],
        ['HU-11', 'RF-13, RF-44']
    ]
    
    for i, traz in enumerate(trazabilidad, start=1):
        row_cells = table_traz.rows[i].cells
        for j, dato in enumerate(traz):
            row_cells[j].text = dato
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
    
    doc.add_paragraph()
    
    # Conclusiones parciales del Epígrafe VI
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi6 = """La ingeniería de requisitos permitió identificar y documentar de manera sistemática las necesidades funcionales y no funcionales del sistema de gestión de pedidos. Los 44 requisitos funcionales, organizados en ocho módulos, cubren todas las operaciones críticas del negocio: autenticación, usuarios, clientes, productos, pedidos, pagos, reportes, auditoría y devoluciones. Los 33 requisitos no funcionales, clasificados según la taxonomía de Sommerville, garantizan que el sistema cumple con estándares de rendimiento, seguridad, usabilidad, compatibilidad, mantenibilidad, fiabilidad, escalabilidad y legalidad. Las 11 historias de usuario establecieron un lenguaje común entre desarrolladores y stakeholders, facilitando la priorización y planificación de iteraciones en la metodología XP. La matriz de trazabilidad HU ↔ RF asegura que cada funcionalidad implementada responde a una necesidad del usuario y que todos los requisitos están cubiertos, garantizando la completitud y coherencia del sistema desarrollado."""
    
    doc.add_paragraph(conclusiones_epi6)
    
    # Salto de página
    doc.add_page_break()

def main():
    filename = 'Informe_Tecnico_PID_NUEVO_V1.docx'
    
    print(f"📄 Abriendo documento existente: {filename}")
    doc = Document(filename)
    
    print("➕ Agregando Epígrafe VI al final del documento...")
    epigrafe_vi(doc)
    
    print("💾 Guardando cambios...")
    doc.save(filename)
    
    print(f"\n✅ Epígrafe VI agregado exitosamente a: {filename}")
    print("📝 Tus correcciones manuales anteriores se han preservado.")

if __name__ == "__main__":
    main()
