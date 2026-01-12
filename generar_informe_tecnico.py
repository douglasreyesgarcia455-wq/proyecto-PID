"""
Generador del Informe Técnico de PID - Sistema de Gestión de Pedidos
Basado en la plantilla UCI y adaptado con información del proyecto
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def crear_informe():
    doc = Document()
    
    # Configurar estilos generales
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # ==================== PORTADA ====================
    portada = doc.add_paragraph()
    portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run = portada.add_run("UNIVERSIDAD DE LAS CIENCIAS INFORMÁTICAS\n")
    run.font.size = Pt(14)
    run.bold = True
    
    run = portada.add_run("FACULTAD 3\n\n\n")
    run.font.size = Pt(13)
    run.bold = True
    
    run = portada.add_run("Sistema de Gestión de Pedidos con Trazabilidad\n\n")
    run.font.size = Pt(16)
    run.bold = True
    
    run = portada.add_run("Informe Técnico de la asignatura de\nProyecto de Investigación y Desarrollo III\n\n\n")
    run.font.size = Pt(13)
    
    run = portada.add_run("Autor: Douglas [Apellidos]\n\n")
    run.font.size = Pt(12)
    
    run = portada.add_run("Tutor: [Nombre del Tutor]\n\n\n")
    run.font.size = Pt(12)
    
    run = portada.add_run(f"La Habana, Enero de 2026\n")
    run.font.size = Pt(12)
    
    run = portada.add_run("Año 67 de la Revolución")
    run.font.size = Pt(11)
    run.italic = True
    
    doc.add_page_break()
    
    # ==================== RESUMEN ====================
    titulo = doc.add_heading("RESUMEN", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    resumen = doc.add_paragraph()
    resumen.add_run(
        "La gestión de pedidos en pequeñas y medianas empresas (PYMES) en Cuba se realiza "
        "actualmente mediante métodos manuales o sistemas no integrados, generando problemas "
        "de trazabilidad, control de inventario inadecuado y falta de auditoría de operaciones. "
        "Este trabajo desarrolló un sistema web para automatizar la gestión integral de pedidos, "
        "permitiendo a usuarios autorizados según roles (Administrador, Supervisor, Vendedor) "
        "registrar clientes, productos, pedidos y pagos con seguimiento completo de estados. "
        "La solución fue implementada utilizando arquitectura de capas con FastAPI como framework "
        "backend, PostgreSQL como sistema de gestión de base de datos y autenticación JWT con "
        "control de acceso basado en roles (RBAC). El sistema incluye 44 requisitos funcionales "
        "documentados mediante diagramas UML de casos de uso y colaboración, garantizando "
        "trazabilidad completa mediante registro de auditoría y gestión de estados de pedidos "
        "(pendiente, pagado, completado, devuelto). Las pruebas funcionales demostraron el "
        "cumplimiento de los requisitos establecidos, validando la solución como herramienta "
        "eficiente para optimizar la gestión comercial, reducir errores operativos y mejorar "
        "la toma de decisiones mediante reportes estadísticos."
    )
    
    doc.add_paragraph()
    palabras = doc.add_paragraph()
    palabras.add_run("PALABRAS CLAVE\n").bold = True
    palabras.add_run(
        "gestión de pedidos; trazabilidad; FastAPI; PostgreSQL; RBAC; sistema web; auditoría"
    )
    
    doc.add_page_break()
    
    # ==================== ABSTRACT ====================
    titulo_abstract = doc.add_heading("ABSTRACT", level=1)
    titulo_abstract.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    abstract = doc.add_paragraph()
    run_abstract = abstract.add_run(
        "Order management in small and medium-sized enterprises (SMEs) in Cuba is currently "
        "carried out through manual methods or non-integrated systems, generating traceability "
        "problems, inadequate inventory control and lack of operations audit. This work developed "
        "a web system to automate comprehensive order management, allowing authorized users "
        "according to roles (Administrator, Supervisor, Salesperson) to register customers, "
        "products, orders and payments with complete status tracking. The solution was implemented "
        "using layered architecture with FastAPI as backend framework, PostgreSQL as database "
        "management system and JWT authentication with role-based access control (RBAC). The "
        "system includes 44 functional requirements documented through UML use case and collaboration "
        "diagrams, ensuring complete traceability through audit logging and order state management "
        "(pending, paid, completed, returned). Functional tests demonstrated compliance with "
        "established requirements, validating the solution as an efficient tool to optimize "
        "commercial management, reduce operational errors and improve decision-making through "
        "statistical reports."
    )
    run_abstract.italic = True
    
    doc.add_paragraph()
    keywords = doc.add_paragraph()
    run_kw = keywords.add_run("KEYWORDS\n")
    run_kw.bold = True
    run_kw.italic = True
    run_kw2 = keywords.add_run(
        "order management; traceability; FastAPI; PostgreSQL; RBAC; web system; audit"
    )
    run_kw2.italic = True
    
    doc.add_page_break()
    
    # ==================== INTRODUCCIÓN ====================
    doc.add_heading("INTRODUCCIÓN", level=1)
    
    doc.add_paragraph(
        "La gestión eficiente de pedidos constituye un proceso fundamental para la operación "
        "de pequeñas y medianas empresas comerciales. En el contexto cubano, muchas PYMES "
        "enfrentan desafíos significativos relacionados con el control manual de inventarios, "
        "la falta de trazabilidad en las transacciones comerciales y la ausencia de mecanismos "
        "automatizados para el seguimiento de pagos y estados de pedidos. Esta situación genera "
        "pérdidas económicas, errores operativos y dificultades para la toma de decisiones "
        "basada en información confiable."
    )
    
    doc.add_paragraph(
        "Actualmente, el proceso de gestión se realiza mediante herramientas genéricas como "
        "hojas de cálculo o cuadernos físicos, donde el registro de clientes, productos, "
        "pedidos y pagos carece de integración, validación automática y mecanismos de auditoría. "
        "Esta metodología manual presenta limitaciones críticas: imposibilidad de consultar el "
        "estado de pedidos en tiempo real, falta de control sobre el stock de productos, ausencia "
        "de registro histórico de operaciones y vulnerabilidad ante errores humanos en cálculos "
        "de montos y validaciones de inventario."
    )
    
    doc.add_paragraph(
        "Esta problemática justifica el desarrollo de una solución tecnológica que centralice "
        "y automatice estos procesos. Por lo tanto, este trabajo se enfoca en el diseño, "
        "desarrollo e implementación de un sistema web que permita la gestión integral de "
        "pedidos con trazabilidad completa. El sistema propuesto facilitará a usuarios con "
        "roles diferenciados (Administrador, Supervisor, Vendedor) el registro y seguimiento "
        "de clientes, productos, inventario, pedidos, pagos y devoluciones, optimizando así "
        "la gestión comercial y mejorando la experiencia operativa de la empresa."
    )
    
    # Problema científico
    doc.add_heading("Problema Científico a Resolver", level=2)
    doc.add_paragraph(
        "¿Cómo diseñar e implementar una solución informática confiable, basada en una "
        "plataforma web y mediante los principios de una ingeniería de software rigurosa, "
        "que permita automatizar la gestión integral de pedidos con trazabilidad completa, "
        "control de inventario en tiempo real, gestión de pagos acumulativos y auditoría "
        "de operaciones, superando las limitaciones de los métodos manuales actuales para "
        "pequeñas y medianas empresas?"
    )
    
    # Objetivo General
    doc.add_heading("Objetivo General", level=2)
    doc.add_paragraph(
        "Diseñar, desarrollar e implementar un sistema web para la gestión integral de "
        "pedidos con trazabilidad completa, utilizando arquitectura de capas, framework "
        "FastAPI, base de datos PostgreSQL y control de acceso basado en roles, que "
        "garantice la automatización de procesos comerciales, el control de inventario "
        "y la generación de reportes estadísticos para la toma de decisiones."
    )
    
    # Objetivos Específicos
    doc.add_heading("Objetivos Específicos", level=2)
    doc.add_paragraph(
        "Los objetivos específicos que guían la investigación y el desarrollo son:"
    )
    
    objetivos = [
        "Analizar el proceso actual de gestión de pedidos en PYMES, identificando los "
        "requisitos funcionales y no funcionales del sistema mediante ingeniería de requisitos.",
        
        "Diseñar la arquitectura de software en capas y modelar la base de datos relacional "
        "que soporte la gestión de usuarios, clientes, productos, inventario, pedidos, pagos "
        "y auditoría de operaciones.",
        
        "Desarrollar los módulos de autenticación (JWT con RBAC), gestión de entidades "
        "(usuarios, clientes, productos), gestión de operaciones (pedidos, pagos, devoluciones) "
        "y reportes estadísticos mediante API REST.",
        
        "Validar el funcionamiento del sistema mediante pruebas unitarias, de integración y "
        "funcionales que demuestren su eficacia en la mejora del proceso de gestión comercial "
        "y el cumplimiento de los 44 requisitos funcionales especificados."
    ]
    
    for i, obj in enumerate(objetivos, 1):
        p = doc.add_paragraph(obj, style='List Number')
    
    # Tareas de Investigación
    doc.add_heading("Tareas de Investigación", level=2)
    doc.add_paragraph(
        "Para alcanzar los objetivos propuestos, se definieron las siguientes tareas:"
    )
    
    tareas = [
        "Realizar un estudio de sistemas de gestión de pedidos existentes y analizar el "
        "proceso actual basado en métodos manuales, para identificar requisitos y "
        "tecnologías aplicables al contexto de PYMES cubanas.",
        
        "Modelar el sistema mediante diagramas UML (casos de uso, colaboración y entidad-relación) "
        "que definan la interacción entre actores, procesos de negocio y estructura de datos.",
        
        "Diseñar y desarrollar la base de datos PostgreSQL que soporte el almacenamiento "
        "transaccional de pedidos, el control de inventario con validaciones y el registro "
        "de auditoría de todas las operaciones.",
        
        "Implementar mediante arquitectura de capas los módulos principales: autenticación "
        "JWT con RBAC, endpoints REST para gestión CRUD de entidades, validación de stock "
        "en pedidos, cálculo automático de estados de pago y generación de reportes.",
        
        "Validar el sistema mediante pruebas técnicas de funcionalidad, seguridad y "
        "rendimiento, ejecutando los 44 casos de prueba correspondientes a los requisitos "
        "funcionales especificados."
    ]
    
    for i, tarea in enumerate(tareas, 1):
        p = doc.add_paragraph(tarea, style='List Number')
    
    # Métodos de Investigación
    doc.add_heading("Métodos de Investigación", level=2)
    
    doc.add_paragraph(
        "El desarrollo de este trabajo se sustenta en un enfoque metodológico mixto que "
        "integra métodos teóricos, empíricos y herramientas especializadas de la Ingeniería "
        "de Software."
    )
    
    doc.add_heading("Métodos Teóricos", level=3)
    
    metodos_teoricos = [
        ("Análisis-Síntesis", "Se empleó durante la revisión sistemática del estado del arte "
         "para descomponer y recomponer los fundamentos de sistemas transaccionales, arquitecturas "
         "de capas, patrones de diseño REST API y mecanismos RBAC."),
        
        ("Inductivo-Deductivo", "Permitió derivar los requisitos específicos del sistema a "
         "partir de la observación del proceso actual manual, generalizando luego en un modelo "
         "aplicable a diferentes tipos de PYMES comerciales."),
        
        ("Método de Modelado", "Se aplicó en el diseño de diagramas UML para representar "
         "requisitos (casos de uso), interacciones (diagramas de colaboración) y estructura "
         "de datos (modelo entidad-relación).")
    ]
    
    for metodo, desc in metodos_teoricos:
        p = doc.add_paragraph()
        p.add_run(f"{metodo}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Métodos Empíricos", level=3)
    
    metodos_empiricos = [
        ("Ingeniería de Requisitos", "Mediante entrevistas con administradores y vendedores "
         "de PYMES, se capturaron las necesidades específicas de trazabilidad, control de "
         "stock y auditoría de operaciones, resultando en 44 requisitos funcionales y 33 "
         "requisitos no funcionales documentados."),
        
        ("Inspección de Base de Datos", "Se analizó la estructura existente de PostgreSQL "
         "mediante script Python (inspect_db.py) para garantizar que los modelos SQLAlchemy "
         "reflejen fielmente el esquema relacional implementado."),
        
        ("Pruebas Funcionales", "Se ejecutaron pruebas manuales y automatizadas para cada "
         "requisito funcional, validando endpoints REST mediante herramientas como Postman "
         "y scripts de prueba Python.")
    ]
    
    for metodo, desc in metodos_empiricos:
        p = doc.add_paragraph()
        p.add_run(f"{metodo}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("Herramientas de Ingeniería de Software", level=3)
    
    herramientas = [
        ("Modelado UML", "Se emplearon diagramas de casos de uso para requisitos, diagramas "
         "de colaboración para flujos de interacción y modelo entidad-relación para diseño "
         "de base de datos, utilizando Draw.io como herramienta CASE."),
        
        ("Arquitectura en Capas", "Se diseñó una estructura modular con separación clara: "
         "capa de presentación (routes), capa de lógica de negocio (services), capa de "
         "acceso a datos (repositories/models) y capa de persistencia (PostgreSQL)."),
        
        ("Arquitectura RESTful", "Se diseñaron APIs específicas siguiendo principios REST "
         "para operaciones CRUD, con respuestas JSON estandarizadas y códigos HTTP semánticos."),
        
        ("Control de Versiones", "Uso de Git/GitHub para versionado de código, permitiendo "
         "trazabilidad de cambios y colaboración en el desarrollo.")
    ]
    
    for herramienta, desc in herramientas:
        p = doc.add_paragraph()
        p.add_run(f"{herramienta}: ").bold = True
        p.add_run(desc)
    
    doc.add_paragraph(
        "Esta integración metodológica permitió abordar tanto la complejidad técnica del "
        "sistema transaccional como las necesidades específicas de trazabilidad y seguridad "
        "identificadas en el contexto de PYMES."
    )
    
    doc.add_page_break()
    
    # ==================== CAPÍTULO I ====================
    doc.add_heading("CAPÍTULO I. ESTUDIO DEL ESTADO DEL ARTE Y FUNDAMENTOS DE LA SOLUCIÓN", level=1)
    
    doc.add_heading("Introducción", level=2)
    doc.add_paragraph(
        "Este capítulo establece los fundamentos teóricos y tecnológicos para el desarrollo "
        "del sistema de gestión de pedidos con trazabilidad. Se analizan conceptos clave de "
        "sistemas transaccionales, arquitecturas de software y tecnologías web modernas. Se "
        "examinan soluciones análogas en el mercado y se justifican las decisiones de ingeniería "
        "de software adoptadas para el proyecto. La metodología integra revisión documental, "
        "análisis comparativo de plataformas existentes y estudio de estándares de la industria, "
        "proporcionando el sustento necesario para el diseño e implementación de la solución "
        "propuesta."
    )
    
    doc.add_heading("1.1 Conceptos Asociados al Tema", level=2)
    
    conceptos = [
        ("Gestión de Pedidos", 
         "Proceso integral que abarca la recepción, validación, procesamiento y seguimiento "
         "de solicitudes de compra realizadas por clientes. Incluye la verificación de "
         "disponibilidad de inventario, cálculo de montos, registro de pagos y actualización "
         "de estados (pendiente, pagado, completado, devuelto). En sistemas informatizados, "
         "este proceso requiere mecanismos transaccionales que garanticen la consistencia de "
         "datos y la trazabilidad de operaciones (Chen et al., 2022)."),
        
        ("Trazabilidad", 
         "Capacidad de un sistema para registrar y rastrear el historial completo de una "
         "operación, identificando quién realizó cada acción, cuándo se ejecutó y qué cambios "
         "se produjeron en los datos. En el contexto de gestión comercial, la trazabilidad "
         "permite auditar pedidos, pagos y modificaciones de inventario, proporcionando "
         "transparencia y facilitando la detección de errores o fraudes (ISO 9001:2015)."),
        
        ("Control de Acceso Basado en Roles (RBAC)", 
         "Modelo de seguridad que restringe el acceso a recursos del sistema según el rol "
         "asignado a cada usuario. En sistemas de gestión de pedidos, permite diferenciar "
         "permisos entre administradores (acceso total), supervisores (gestión de inventario "
         "y pedidos) y vendedores (solo creación de pedidos), garantizando segregación de "
         "funciones y prevención de accesos no autorizados (Ferraiolo et al., 2001)."),
        
        ("API REST", 
         "Interfaz de programación de aplicaciones que utiliza el protocolo HTTP y principios "
         "de arquitectura REST (Representational State Transfer) para exponer operaciones del "
         "sistema mediante endpoints con métodos estándar (GET, POST, PUT, DELETE). Facilita "
         "la integración con aplicaciones frontend, móviles o sistemas externos, garantizando "
         "interoperabilidad y escalabilidad (Fielding, 2000)."),
        
        ("Transaccionalidad ACID", 
         "Conjunto de propiedades que garantizan la confiabilidad de operaciones en bases de "
         "datos: Atomicidad (todo o nada), Consistencia (reglas de integridad respetadas), "
         "Aislamiento (operaciones concurrentes no interfieren) y Durabilidad (cambios "
         "confirmados persisten). Esencial en sistemas de gestión de pedidos para evitar "
         "inconsistencias en inventario y registros de pago (Gray & Reuter, 1992).")
    ]
    
    for concepto, definicion in conceptos:
        p = doc.add_paragraph()
        p.add_run(f"{concepto}: ").bold = True
        p.add_run(definicion)
    
    doc.add_heading("1.2 Análisis de Soluciones Existentes", level=2)
    
    doc.add_paragraph(
        "El estudio de soluciones existentes revela diversas plataformas de gestión comercial "
        "con diferentes enfoques tecnológicos y funcionales. A continuación se presenta un "
        "análisis comparativo de sistemas relevantes:"
    )
    
    # Crear tabla de análisis comparativo
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers = ['Sistema', 'Tipo', 'Trazabilidad', 'Control RBAC', 'Tecnología Backend']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Datos
    datos = [
        ['Odoo', 'ERP Open Source', 'Completa', 'Sí', 'Python (Flask)'],
        ['ERPNext', 'ERP Open Source', 'Completa', 'Sí', 'Python (Frappe)'],
        ['Shopify', 'SaaS E-commerce', 'Parcial', 'Básico', 'Ruby on Rails'],
        ['WooCommerce', 'Plugin WordPress', 'Básica', 'No', 'PHP'],
        ['Solución Propuesta', 'Sistema ad-hoc', 'Completa', 'Sí', 'Python (FastAPI)']
    ]
    
    for i, row_data in enumerate(datos, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "El análisis identifica como puntos coincidentes la gestión de datos comerciales y "
        "el control de inventario. Sin embargo, se evidencia una brecha en soluciones "
        "específicas para PYMES cubanas que combinen simplicidad de uso, arquitectura moderna "
        "(FastAPI), trazabilidad completa mediante auditoría y control granular de acceso RBAC "
        "sin la complejidad de sistemas ERP empresariales completos."
    )
    
    doc.add_heading("1.3 Fundamentación del Proceso de Software a Desarrollar", level=2)
    
    doc.add_paragraph(
        "El dominio de aplicación requiere alta confiabilidad en datos transaccionales, "
        "validaciones de integridad referencial, capacidad de respuesta inmediata para "
        "operaciones CRUD y escalabilidad para múltiples usuarios concurrentes. El proceso "
        "debe garantizar seguridad en autenticación, autorización granular y disponibilidad "
        "continua del sistema."
    )
    
    doc.add_heading("1.3.1 Enfoque de Ingeniería de Software", level=3)
    
    doc.add_paragraph(
        "Se adopta un enfoque de desarrollo iterativo e incremental basado en los siguientes "
        "criterios del proyecto:"
    )
    
    criterios = [
        "Equipo: 1 desarrollador con conocimiento en Python y bases de datos relacionales",
        "Requisitos: Bien definidos (44 RF + 33 RNF documentados) con alta estabilidad",
        "Criticidad: Alta (sistema transaccional con manejo de datos comerciales sensibles)",
        "Tecnología: Precedencia en FastAPI, PostgreSQL y arquitectura de capas",
        "Plazo: 4 meses para desarrollo, pruebas y documentación"
    ]
    
    for criterio in criterios:
        doc.add_paragraph(criterio, style='List Bullet')
    
    doc.add_paragraph(
        "Esta combinación favorece un modelo de desarrollo en capas con entregas incrementales "
        "por módulos funcionales, priorizando autenticación → gestión de entidades → "
        "operaciones transaccionales → reportes."
    )
    
    doc.add_heading("1.3.2 Arquitectura de Software Seleccionada", level=3)
    
    doc.add_paragraph(
        "Se implementa una arquitectura en capas (Layered Architecture) que separa "
        "responsabilidades y facilita el mantenimiento:"
    )
    
    capas = [
        ("Capa de Presentación (Routes)", 
         "Endpoints REST que exponen la funcionalidad del sistema mediante API HTTP. "
         "Maneja validación de entrada, serialización JSON y códigos de respuesta HTTP."),
        
        ("Capa de Lógica de Negocio (Services)", 
         "Implementa las reglas del negocio: validación de stock antes de crear pedidos, "
         "cálculo de estados de pago, actualización de inventario tras ventas, registro "
         "de auditoría de operaciones."),
        
        ("Capa de Acceso a Datos (Repositories/Models)", 
         "Modelos SQLAlchemy que mapean entidades a tablas PostgreSQL y encapsulan "
         "operaciones de persistencia (CRUD básico)."),
        
        ("Capa de Persistencia (PostgreSQL)", 
         "Base de datos relacional que garantiza transaccionalidad ACID, integridad "
         "referencial y consultas eficientes mediante índices.")
    ]
    
    for capa, desc in capas:
        p = doc.add_paragraph()
        p.add_run(f"{capa}: ").bold = True
        p.add_run(desc)
    
    doc.add_heading("1.4 Herramientas y Tecnologías", level=2)
    
    doc.add_heading("1.4.1 Herramienta CASE", level=3)
    doc.add_paragraph(
        "Draw.io (diagrams.net) para modelado de diagramas UML (casos de uso, colaboración) "
        "y diseño de arquitectura, seleccionada por su gratuidad, exportación a múltiples "
        "formatos y facilidad de uso para documentación técnica."
    )
    
    doc.add_heading("1.4.2 Lenguaje de Modelado", level=3)
    doc.add_paragraph(
        "UML (Unified Modeling Language) versión 2.5, estandarizado para representación de "
        "sistemas mediante diagramas estructurales (casos de uso) y comportamentales "
        "(colaboración, secuencia)."
    )
    
    doc.add_heading("1.4.3 Framework Backend", level=3)
    doc.add_paragraph(
        "FastAPI 0.104+ como framework web moderno para Python, seleccionado por las "
        "siguientes ventajas:"
    )
    
    ventajas_fastapi = [
        "Alto rendimiento: comparable a Node.js y Go gracias a Starlette y Pydantic",
        "Tipado estático: validación automática de datos mediante type hints de Python",
        "Documentación automática: generación de OpenAPI/Swagger sin configuración adicional",
        "Asincronía nativa: soporte async/await para operaciones I/O eficientes",
        "Desarrollo rápido: reducción de código repetitivo (boilerplate)"
    ]
    
    for ventaja in ventajas_fastapi:
        doc.add_paragraph(ventaja, style='List Bullet')
    
    doc.add_heading("1.4.4 ORM (Object-Relational Mapping)", level=3)
    doc.add_paragraph(
        "SQLAlchemy 2.0+ como biblioteca de mapeo objeto-relacional, proporcionando:"
    )
    
    ventajas_sqlalchemy = [
        "Abstracción de base de datos: independencia del motor SQL utilizado",
        "Modelos declarativos: definición clara de entidades y relaciones",
        "Migraciones: gestión de cambios de esquema mediante Alembic",
        "Consultas tipadas: prevención de errores mediante type hints"
    ]
    
    for ventaja in ventajas_sqlalchemy:
        doc.add_paragraph(ventaja, style='List Bullet')
    
    doc.add_heading("1.4.5 Sistema de Gestión de Base de Datos", level=3)
    doc.add_paragraph(
        "PostgreSQL 16 como SGBD relacional, justificado por:"
    )
    
    ventajas_postgres = [
        "Transaccionalidad ACID completa: garantiza consistencia en operaciones críticas",
        "Integridad referencial: claves foráneas y restricciones para validación de datos",
        "Rendimiento: índices B-tree, BRIN, GIN para consultas eficientes",
        "Tipos de datos avanzados: JSON, arrays, tipos personalizados",
        "Licencia: Open Source (PostgreSQL License), sin costos de licenciamiento",
        "Madurez: más de 30 años de desarrollo activo"
    ]
    
    for ventaja in ventajas_postgres:
        doc.add_paragraph(ventaja, style='List Bullet')
    
    doc.add_heading("1.4.6 Autenticación y Autorización", level=3)
    doc.add_paragraph(
        "JSON Web Tokens (JWT) con algoritmo HS256 para autenticación stateless, combinado "
        "con RBAC (Role-Based Access Control) para autorización granular. Se implementan "
        "tres roles:"
    )
    
    roles = [
        ("Administrador", "Acceso total: gestión de usuarios, roles, respaldos, configuración"),
        ("Supervisor", "Gestión de inventario, pedidos, pagos, reportes (sin gestión de usuarios)"),
        ("Vendedor", "Solo creación de pedidos y registro de pagos (lectura de productos/clientes)")
    ]
    
    for rol, permisos in roles:
        p = doc.add_paragraph()
        p.add_run(f"{rol}: ").bold = True
        p.add_run(permisos)
    
    doc.add_heading("1.4.7 Entorno de Desarrollo", level=3)
    doc.add_paragraph(
        "Visual Studio Code como IDE principal, con extensiones para Python (Pylance), "
        "Git (GitLens), y herramientas de formato (Black, Ruff). Control de versiones "
        "mediante Git/GitHub para trazabilidad de cambios."
    )
    
    doc.add_heading("1.4.8 Lenguaje de Programación", level=3)
    doc.add_paragraph(
        "Python 3.11+ como lenguaje principal, aprovechando:"
    )
    
    ventajas_python = [
        "Ecosistema maduro: librerías robustas para desarrollo web (FastAPI, SQLAlchemy)",
        "Legibilidad: sintaxis clara que facilita mantenimiento",
        "Tipado gradual: type hints para detección temprana de errores",
        "Comunidad activa: documentación extensa y soporte comunitario"
    ]
    
    for ventaja in ventajas_python:
        doc.add_paragraph(ventaja, style='List Bullet')
    
    doc.add_heading("Conclusiones Parciales", level=2)
    
    doc.add_paragraph(
        "El estudio del estado del arte permitió identificar los conceptos fundamentales "
        "que sustentan el desarrollo de sistemas transaccionales de gestión comercial: "
        "trazabilidad, RBAC, APIs REST y arquitecturas en capas. El análisis comparativo "
        "de soluciones existentes evidenció la necesidad específica de sistemas adaptados "
        "a PYMES que combinen simplicidad operativa con robustez técnica."
    )
    
    doc.add_paragraph(
        "Las herramientas y tecnologías seleccionadas (FastAPI, PostgreSQL, SQLAlchemy, JWT) "
        "conforman un stack moderno y coherente con los requisitos identificados, garantizando "
        "rendimiento, escalabilidad y mantenibilidad del sistema. La arquitectura en capas "
        "propuesta facilita la separación de responsabilidades y el desarrollo incremental "
        "por módulos funcionales."
    )
    
    doc.add_paragraph(
        "El capítulo establece las bases técnicas necesarias para proceder con el diseño "
        "detallado del modelo de negocio, la especificación de requisitos y la implementación "
        "de la solución propuesta, asegurando su alineación con las mejores prácticas actuales "
        "de ingeniería de software."
    )
    
    doc.add_page_break()
    
    # ==================== CAPÍTULO II ====================
    doc.add_heading("CAPÍTULO II. MODELADO DEL CONTEXTO Y ESPECIFICACIÓN DE REQUISITOS", level=1)
    
    doc.add_heading("Introducción", level=2)
    doc.add_paragraph(
        "El presente capítulo tiene como objetivo fundamental analizar el modelo de negocio "
        "actual para la gestión de pedidos en PYMES y proponer una solución informática que "
        "optimice sus procesos. Mediante el empleo de técnicas de modelado empresarial y de "
        "requisitos, se caracteriza el sistema existente identificando sus reglas de operación, "
        "actores involucrados y flujos de trabajo. Posteriormente, se establece el modelo "
        "mejorado que sustenta el desarrollo del sistema, especificando los nuevos procesos, "
        "actores, reglas de negocio y requisitos funcionales y no funcionales. La metodología "
        "de análisis se apoya en diagramas UML y especificaciones estructuradas que garantizan "
        "una transición coherente entre el contexto actual y la propuesta de valor del software "
        "a desarrollar."
    )
    
    doc.add_heading("2.1 Reglas del Negocio", level=2)
    
    doc.add_paragraph(
        "A continuación se especifican las reglas que rigen el proceso de gestión de pedidos:"
    )
    
    # Crear tabla de reglas de negocio
    table_reglas = doc.add_table(rows=11, cols=4)
    table_reglas.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers_reglas = ['No', 'Clasificación', 'Nombre', 'Descripción']
    for i, header in enumerate(headers_reglas):
        cell = table_reglas.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Datos
    reglas_datos = [
        ['1', 'Hecho', 'Autenticación Obligatoria', 
         'Todo usuario debe autenticarse mediante credenciales válidas antes de acceder al sistema.'],
        ['2', 'Restricción', 'Control de Acceso por Roles', 
         'Las operaciones disponibles para cada usuario están determinadas por su rol asignado (Admin, Supervisor, Vendedor).'],
        ['3', 'Hecho', 'Registro de Clientes', 
         'Los clientes deben registrarse en el sistema antes de realizar pedidos, incluyendo nombre, teléfono y dirección.'],
        ['4', 'Restricción', 'Validación de Stock', 
         'No se puede crear un pedido si algún producto solicitado tiene cantidad insuficiente en inventario.'],
        ['5', 'Computacional', 'Actualización Automática de Stock', 
         'Al confirmar un pedido, el sistema reduce automáticamente las cantidades de productos del inventario.'],
        ['6', 'Hecho', 'Estado Inicial de Pedido', 
         'Todo pedido nuevo se crea con estado "pendiente" hasta que se registren pagos.'],
        ['7', 'Computacional', 'Cálculo de Estado de Pago', 
         'Si la suma de pagos registrados es mayor o igual al total del pedido, el estado cambia automáticamente a "pagado".'],
        ['8', 'Restricción', 'Pagos Acumulativos', 
         'Se permiten múltiples pagos parciales para un mismo pedido hasta cubrir el monto total.'],
        ['9', 'Restricción', 'Devoluciones Condicionadas', 
         'Solo pedidos en estado "pagado" pueden ser devueltos, restaurando el inventario correspondiente.'],
        ['10', 'Facilitador', 'Auditoría Automática', 
         'Toda operación de creación, modificación o eliminación se registra automáticamente con usuario, fecha y acción realizada.']
    ]
    
    for i, row_data in enumerate(reglas_datos, 1):
        for j, cell_data in enumerate(row_data):
            table_reglas.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("2.2 Actores del Negocio", level=2)
    
    doc.add_paragraph(
        "Los actores identificados que interactúan con el sistema son:"
    )
    
    # Crear tabla de actores
    table_actores = doc.add_table(rows=5, cols=2)
    table_actores.style = 'Light Grid Accent 1'
    
    # Encabezados
    headers_actores = ['Actor', 'Descripción']
    for i, header in enumerate(headers_actores):
        cell = table_actores.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Datos
    actores_datos = [
        ['Administrador', 
         'Usuario con privilegios totales: gestiona usuarios, asigna roles, realiza respaldos, configura el sistema y accede a todos los módulos.'],
        ['Supervisor', 
         'Usuario con permisos de gestión: administra productos e inventario, crea y modifica pedidos, registra pagos, genera reportes. No puede gestionar usuarios.'],
        ['Vendedor', 
         'Usuario con permisos limitados: solo puede crear pedidos, registrar pagos y consultar información de productos y clientes. Sin acceso a inventario ni reportes.'],
        ['Cliente', 
         'Entidad externa que realiza compras. Sus datos se registran en el sistema para asociarlos a pedidos y generar historial de transacciones.']
    ]
    
    for i, row_data in enumerate(actores_datos, 1):
        for j, cell_data in enumerate(row_data):
            table_actores.rows[i].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    doc.add_heading("2.3 Especificación de Requisitos Funcionales", level=2)
    
    doc.add_paragraph(
        "El sistema contempla 44 requisitos funcionales organizados por módulos. A continuación "
        "se presenta un resumen agrupado por categorías funcionales. Los diagramas de casos de "
        "uso y diagramas de colaboración detallados para cada requisito se encuentran en los "
        "Anexos."
    )
    
    doc.add_heading("2.3.1 Gestión de Usuarios (RF01-RF08)", level=3)
    
    rf_usuarios = [
        "RF01: El sistema debe permitir crear usuarios internos con credenciales únicas",
        "RF02: El sistema debe permitir modificar datos de usuarios existentes",
        "RF03: El sistema debe permitir eliminar o desactivar usuarios",
        "RF04: El sistema debe permitir listar usuarios con filtros por rol",
        "RF05: El sistema debe permitir consultar detalles de un usuario específico",
        "RF06: El sistema debe permitir al usuario consultar su propio perfil",
        "RF07: El sistema debe permitir iniciar sesión mediante usuario y contraseña",
        "RF08: El sistema debe permitir cerrar sesión invalidando el token JWT"
    ]
    
    for rf in rf_usuarios:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.2 Gestión de Clientes (RF09-RF13)", level=3)
    
    rf_clientes = [
        "RF09: El sistema debe permitir crear clientes con datos básicos (nombre, teléfono, dirección)",
        "RF10: El sistema debe permitir modificar información de clientes",
        "RF11: El sistema debe permitir eliminar clientes sin pedidos asociados",
        "RF12: El sistema debe permitir listar clientes con paginación",
        "RF13: El sistema debe permitir consultar detalles de un cliente y su historial de pedidos"
    ]
    
    for rf in rf_clientes:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.3 Gestión de Productos e Inventario (RF14-RF20)", level=3)
    
    rf_productos = [
        "RF14: El sistema debe permitir crear productos con nombre, descripción, precio, cantidad",
        "RF15: El sistema debe permitir modificar datos de productos",
        "RF16: El sistema debe permitir eliminar productos sin pedidos asociados",
        "RF17: El sistema debe permitir listar productos con filtros y búsqueda",
        "RF18: El sistema debe proporcionar un catálogo público de productos disponibles",
        "RF19: El sistema debe reducir automáticamente el stock al confirmar pedidos",
        "RF20: El sistema debe alertar cuando un producto alcance stock mínimo"
    ]
    
    for rf in rf_productos:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.4 Gestión de Pedidos (RF21-RF27)", level=3)
    
    rf_pedidos = [
        "RF21: El sistema debe permitir crear pedidos con múltiples productos",
        "RF22: El sistema debe validar disponibilidad de stock antes de crear pedidos",
        "RF23: El sistema debe permitir listar pedidos con filtros por estado",
        "RF24: El sistema debe permitir consultar detalles completos de un pedido",
        "RF25: El sistema debe permitir actualizar pedidos en estado pendiente",
        "RF26: El sistema debe asignar automáticamente estado 'pendiente' a pedidos nuevos",
        "RF27: El sistema debe cambiar automáticamente estado a 'pagado' cuando pagos cubran el total"
    ]
    
    for rf in rf_pedidos:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.5 Gestión de Pagos (RF28-RF33)", level=3)
    
    rf_pagos = [
        "RF28: El sistema debe permitir registrar pagos asociados a pedidos",
        "RF29: El sistema debe validar que el monto de pago no exceda el pendiente",
        "RF30: El sistema debe actualizar automáticamente el total pagado del pedido",
        "RF31: El sistema debe permitir listar pagos con filtros por pedido",
        "RF32: El sistema debe permitir consultar detalles de un pago específico",
        "RF33: El sistema debe proporcionar resumen de pagos por pedido"
    ]
    
    for rf in rf_pagos:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.6 Reportes y Estadísticas (RF34-RF38)", level=3)
    
    rf_reportes = [
        "RF34: El sistema debe generar estadísticas diarias de ventas",
        "RF35: El sistema debe generar estadísticas mensuales de ventas",
        "RF36: El sistema debe proporcionar resumen de pedidos pendientes",
        "RF37: El sistema debe permitir exportar reportes en formato PDF",
        "RF38: El sistema debe permitir exportar reportes en formato Excel"
    ]
    
    for rf in rf_reportes:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_heading("2.3.7 Auditoría y Devoluciones (RF39-RF44)", level=3)
    
    rf_otros = [
        "RF39: El sistema debe permitir listar logs de auditoría",
        "RF40: El sistema debe permitir filtrar logs por usuario, fecha y tipo de acción",
        "RF41: El sistema debe permitir registrar devoluciones de pedidos",
        "RF42: El sistema debe restaurar inventario automáticamente al procesar devoluciones",
        "RF43: El sistema debe permitir consultar detalles de devoluciones",
        "RF44: El sistema debe permitir buscar clientes por criterios múltiples"
    ]
    
    for rf in rf_otros:
        doc.add_paragraph(rf, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        "Nota: Los diagramas de casos de uso y diagramas de colaboración para cada RF se "
        "encuentran en la carpeta diagramas_analisis/ y diagramas_colaboracion/ del proyecto."
    )
    
    doc.add_heading("2.4 Especificación de Requisitos No Funcionales", level=2)
    
    doc.add_paragraph(
        "El sistema contempla 33 requisitos no funcionales categorizados en:"
    )
    
    rnf_categorias = [
        ("Rendimiento (RNF01-RNF04)", [
            "Tiempo de respuesta de consultas < 2 segundos",
            "Soporte para 500 usuarios concurrentes sin degradación",
            "Generación de reportes < 10 segundos",
            "Carga de catálogo < 3 segundos"
        ]),
        
        ("Seguridad (RNF05-RNF10)", [
            "Contraseñas almacenadas con hash bcrypt/argon2",
            "Comunicación mediante HTTPS exclusivamente",
            "Autenticación obligatoria mediante JWT",
            "Registro de auditoría de todas las operaciones",
            "Autorización basada en roles (RBAC) en backend",
            "Recuperación de contraseña mediante correo electrónico"
        ]),
        
        ("Usabilidad (RNF11-RNF15)", [
            "Interfaz intuitiva para usuarios no técnicos",
            "Diseño responsive para móviles, tabletas y PC",
            "Botones y funciones claramente etiquetados",
            "Confirmaciones visuales de acciones (toasts, alertas)",
            "Panel de administración con navegación sencilla"
        ]),
        
        ("Compatibilidad (RNF16-RNF19)", [
            "Compatibilidad con navegadores modernos (Chrome, Firefox, Edge)",
            "Base de datos PostgreSQL (versión 12+)",
            "Exportación en formatos estándar (PDF, XLSX, CSV)",
            "Integración mediante API REST documentada (OpenAPI)"
        ]),
        
        ("Mantenibilidad (RNF20-RNF23)", [
            "Código documentado y versionado en GitHub",
            "Arquitectura modular (capas separadas)",
            "Capacidad de agregar nuevos roles sin modificar estructura base",
            "Actualizaciones sin pérdida de datos"
        ]),
        
        ("Fiabilidad y Disponibilidad (RNF24-RNF27)", [
            "Disponibilidad mínima del 99% mensual",
            "Respaldo automático diario",
            "Recuperación < 10 minutos ante fallas",
            "Prevención de pérdida de datos ante desconexiones"
        ]),
        
        ("Escalabilidad (RNF28-RNF30)", [
            "Capacidad de ampliarse para más productos, usuarios y pedidos",
            "Arquitectura preparada para conexión con apps móviles",
            "Crecimiento sin reestructuración completa"
        ]),
        
        ("Legalidad y Privacidad (RNF31-RNF33)", [
            "Cumplimiento con leyes de protección de datos (GDPR o equivalente)",
            "Datos de clientes no compartidos sin consentimiento",
            "Capacidad de eliminar cuentas y datos personales"
        ])
    ]
    
    for categoria, requisitos in rnf_categorias:
        doc.add_heading(categoria, level=3)
        for req in requisitos:
            doc.add_paragraph(req, style='List Bullet')
    
    doc.add_heading("Conclusiones Parciales", level=2)
    
    doc.add_paragraph(
        "El modelado del contexto permitió identificar las reglas de negocio fundamentales "
        "que rigen la gestión de pedidos en PYMES, así como los actores principales que "
        "interactúan con el sistema (Administrador, Supervisor, Vendedor, Cliente). La "
        "especificación de 44 requisitos funcionales organizados en 7 módulos establece "
        "una base sólida para el diseño e implementación de la solución."
    )
    
    doc.add_paragraph(
        "Los 33 requisitos no funcionales garantizan que el sistema no solo cumpla con las "
        "funcionalidades esperadas, sino que también satisfaga estándares de rendimiento, "
        "seguridad, usabilidad y escalabilidad necesarios para su adopción en entornos "
        "productivos. La combinación de requisitos funcionales documentados mediante diagramas "
        "UML y requisitos no funcionales cuantificables proporciona una especificación completa "
        "para las etapas subsecuentes de diseño, implementación y validación del sistema."
    )
    
    # Guardar documento
    doc.save("Informe_Tecnico_PID_Gestion_Pedidos.docx")
    print("✅ Documento generado: Informe_Tecnico_PID_Gestion_Pedidos.docx")
    print("\n📄 Contenido generado:")
    print("   - Portada completa")
    print("   - Resumen y Abstract")
    print("   - Introducción con diseño metodológico")
    print("   - Capítulo I: Estado del arte y fundamentos")
    print("   - Capítulo II: Modelado del contexto y requisitos")
    print("\n⚠️  Pendiente de completar manualmente:")
    print("   - Capítulo III: Diseño e implementación (con código)")
    print("   - Capítulo IV: Validación y pruebas")
    print("   - Conclusiones")
    print("   - Recomendaciones")
    print("   - Referencias bibliográficas")
    print("   - Anexos")
    print("   - Historias de usuario (Sección 2.3)")

if __name__ == "__main__":
    try:
        crear_informe()
    except Exception as e:
        print(f"❌ Error: {e}")
        import traceback
        traceback.print_exc()
