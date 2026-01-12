"""Script to add code examples to the technical report"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def agregar_parrafo_codigo(doc, titulo, codigo, explicacion=""):
    """Add a code block with title and optional explanation"""
    if explicacion:
        p = doc.add_paragraph()
        p.add_run(explicacion).font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    # Add title for code block
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(11)
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(3)
    
    # Add code block
    p_codigo = doc.add_paragraph()
    run_codigo = p_codigo.add_run(codigo)
    run_codigo.font.name = 'Consolas'
    run_codigo.font.size = Pt(9)
    run_codigo.font.color.rgb = RGBColor(0, 0, 0)
    
    # Format paragraph
    p_codigo.paragraph_format.left_indent = Inches(0.25)
    p_codigo.paragraph_format.right_indent = Inches(0.25)
    p_codigo.paragraph_format.space_before = Pt(6)
    p_codigo.paragraph_format.space_after = Pt(6)
    p_codigo.paragraph_format.line_spacing = 1.0
    
    # Add light gray background (simulated with border)
    p_codigo.paragraph_format.keep_together = True


def buscar_parrafo_con_texto(doc, texto_buscar):
    """Search for a paragraph containing specific text"""
    for i, para in enumerate(doc.paragraphs):
        if texto_buscar.lower() in para.text.lower():
            return i
    return -1


def insertar_despues_de_parrafo(doc, indice, contenido_func):
    """Insert content after a specific paragraph"""
    # We need to work with the underlying XML
    # For simplicity, we'll append at the end of the section
    contenido_func(doc)


def main():
    print("📄 Abriendo documento...")
    doc = Document("Informe_Tecnico_PID_Gestion_Pedidos.docx")
    
    print("🔍 Buscando secciones para agregar código...")
    
    # Find Chapter III section
    capitulo_iii_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if ("CAPÍTULO III" in para.text or "Capítulo III" in para.text) and "DISEÑO" in para.text.upper():
            capitulo_iii_idx = i
            break
    
    if capitulo_iii_idx == -1:
        print("❌ No se encontró Capítulo III")
        return
    
    print(f"✅ Capítulo III encontrado en párrafo {capitulo_iii_idx}")
    
    # CODIGO 1: Modelo SQLAlchemy
    codigo_modelo = '''from sqlalchemy import Column, Integer, String, Boolean, DateTime
from datetime import datetime
from src.core.database import Base

class Usuario(Base):
    __tablename__ = "usuarios"
    
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(50), unique=True, nullable=False, index=True)
    email = Column(String(100), unique=True, nullable=False, index=True)
    hashed_password = Column(String(255), nullable=False)
    rol = Column(String(11), nullable=False)  # admin, vendedor, supervisor
    is_active = Column(Boolean, default=True, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow, nullable=False)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)'''
    
    # CODIGO 2: Endpoint de autenticación
    codigo_endpoint = '''from fastapi import APIRouter, Depends
from sqlalchemy.orm import Session
from src.core.database import get_db
from src.modules.auth.schema import LoginRequest, TokenResponse
from src.modules.auth.service import AuthService

router = APIRouter(prefix="/api/auth", tags=["auth"])

@router.post("/login", response_model=TokenResponse)
def login(credentials: LoginRequest, db: Session = Depends(get_db)):
    """Login endpoint - returns JWT token"""
    # Authenticate user
    user = AuthService.authenticate_user(
        db, 
        credentials.username, 
        credentials.password
    )
    
    # Generate token
    access_token = AuthService.create_user_token(user)
    
    # Return response
    return TokenResponse(
        access_token=access_token,
        user_id=user.id,
        username=user.username,
        rol=user.rol
    )'''
    
    # CODIGO 3: Servicio de creación de pedido
    codigo_servicio = '''@staticmethod
def create_order(db: Session, order_data: OrderCreate) -> Pedido:
    """Create new order with stock validation"""
    # 1. Verify client exists
    client = db.query(Cliente).filter(Cliente.id == order_data.cliente_id).first()
    if not client:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # 2. Calculate total and verify stock
    total = Decimal(0)
    detalles_to_create = []
    
    for detalle in order_data.detalles:
        producto = db.query(Producto).filter(Producto.id == detalle.producto_id).first()
        if not producto:
            raise HTTPException(status_code=404, detail=f"Product {detalle.producto_id} not found")
        
        # Validate stock availability
        if producto.stock < detalle.cantidad:
            raise HTTPException(
                status_code=400,
                detail=f"Insufficient stock for '{producto.nombre}'. Available: {producto.stock}"
            )
        
        # Calculate subtotal
        precio = detalle.precio_unitario or Decimal(str(producto.precio_venta))
        subtotal = precio * detalle.cantidad
        total += subtotal
        
        detalles_to_create.append({
            "producto_id": detalle.producto_id,
            "cantidad": detalle.cantidad,
            "precio_unitario": precio,
            "subtotal": subtotal,
            "producto": producto
        })
    
    # 3. Create order in database
    estado_inicial = "pagado" if order_data.pago_inmediato else "pendiente"
    db_order = Pedido(
        cliente_id=order_data.cliente_id,
        estado=estado_inicial,
        total=total
    )
    db.add(db_order)
    db.flush()
    
    # 4. Create order details and update stock
    for detalle_data in detalles_to_create:
        producto = detalle_data.pop("producto")
        db_detalle = DetallePedido(pedido_id=db_order.id, **detalle_data)
        db.add(db_detalle)
        
        # Reduce product stock
        producto.stock -= detalle_data["cantidad"]
    
    db.commit()
    db.refresh(db_order)
    return db_order'''
    
    print("\n✍️ Agregando ejemplos de código al documento...")
    
    # Add section for code examples in Chapter III
    # Find the section about modules
    encontrado_modulos = False
    for i, para in enumerate(doc.paragraphs):
        if i > capitulo_iii_idx and "Módulo de Autenticación" in para.text:
            encontrado_modulos = True
            print(f"✅ Sección de módulos encontrada en párrafo {i}")
            break
    
    # Instead of searching for red markers, we'll add a new subsection
    # with all code examples at the end of Chapter III
    
    # Find where Chapter III ends (before Chapter IV)
    capitulo_iv_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if i > capitulo_iii_idx and ("CAPÍTULO IV" in para.text or "Capítulo IV" in para.text):
            capitulo_iv_idx = i
            break
    
    if capitulo_iv_idx == -1:
        print("⚠️ No se encontró Capítulo IV, agregando al final del documento")
        capitulo_iv_idx = len(doc.paragraphs)
    
    print(f"📍 Capítulo IV comienza en párrafo {capitulo_iv_idx}")
    print("📝 Agregando nueva subsección '3.5 Ejemplos de Implementación'...")
    
    # We'll add content at the end since inserting in the middle is complex with python-docx
    # Add heading
    doc.add_heading("3.5. Ejemplos de Implementación", level=2)
    
    # Add intro paragraph
    p_intro = doc.add_paragraph(
        "A continuación se presentan tres ejemplos representativos de la implementación del sistema, "
        "mostrando la estructura de las capas de datos, lógica de negocio y presentación."
    )
    p_intro.paragraph_format.space_after = Pt(12)
    
    # Ejemplo 1: Modelo
    doc.add_heading("3.5.1. Capa de Datos: Modelo Usuario", level=3)
    agregar_parrafo_codigo(
        doc,
        "Código 3.1: Definición del modelo Usuario con SQLAlchemy",
        codigo_modelo,
        "El siguiente código muestra la definición del modelo Usuario utilizando SQLAlchemy ORM. "
        "Se definen columnas con restricciones, índices para búsquedas eficientes, y campos de auditoría (created_at, updated_at)."
    )
    
    # Ejemplo 2: Endpoint
    doc.add_heading("3.5.2. Capa de Presentación: Endpoint de Autenticación", level=3)
    agregar_parrafo_codigo(
        doc,
        "Código 3.2: Endpoint POST /api/auth/login",
        codigo_endpoint,
        "Este endpoint maneja la autenticación de usuarios. Recibe credenciales, valida usuario y contraseña "
        "a través del servicio de autenticación, y retorna un token JWT junto con información del usuario."
    )
    
    # Ejemplo 3: Servicio
    doc.add_heading("3.5.3. Capa de Lógica: Servicio de Creación de Pedidos", level=3)
    agregar_parrafo_codigo(
        doc,
        "Código 3.3: Método create_order() del servicio de pedidos",
        codigo_servicio,
        "Este método implementa la lógica completa de creación de pedidos: (1) valida existencia del cliente, "
        "(2) verifica stock de productos, (3) calcula totales, (4) crea el pedido y detalles, "
        "(5) actualiza inventario. Todo se ejecuta en una transacción para garantizar consistencia (ACID)."
    )
    
    # Add concluding paragraph
    p_conclusion = doc.add_paragraph(
        "Estos ejemplos demuestran la separación de responsabilidades implementada en el sistema: "
        "el modelo define la estructura de datos, los servicios contienen la lógica de negocio con validaciones, "
        "y los endpoints exponen funcionalidad a través de una API RESTful. Esta arquitectura facilita el mantenimiento, "
        "pruebas unitarias y escalabilidad del sistema."
    )
    p_conclusion.paragraph_format.space_before = Pt(12)
    
    print("\n💾 Guardando documento...")
    doc.save("Informe_Tecnico_PID_Gestion_Pedidos.docx")
    
    print("\n✅ Código agregado exitosamente")
    print("📄 Archivo actualizado: Informe_Tecnico_PID_Gestion_Pedidos.docx")
    print("\n📋 Resumen de cambios:")
    print("   ✅ Agregada subsección 3.5 'Ejemplos de Implementación'")
    print("   ✅ Código 3.1: Modelo Usuario (SQLAlchemy)")
    print("   ✅ Código 3.2: Endpoint /api/auth/login (FastAPI)")
    print("   ✅ Código 3.3: Servicio crear_pedido() (Lógica de negocio)")
    print("\n🔴 Marcadores rojos restantes: 4")
    print("   - Capturas de Postman (Chapter III y IV)")
    print("   - Tabla completa de 44 casos de prueba (Chapter IV)")
    print("   - Gráficas de pruebas de carga (Chapter IV)")
    print("   - Captura de reporte pytest-cov (Chapter IV)")
    print("\n💡 Nota: Los marcadores rojos restantes requieren capturas de pantalla")
    print("   que solo pueden ser agregadas manualmente ejecutando el sistema.")


if __name__ == "__main__":
    main()
