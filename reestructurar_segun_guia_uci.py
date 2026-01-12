"""Script para reestructurar el documento según la guía UCI oficial"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy


def crear_documento_reestructurado():
    """Crea nuevo documento con estructura correcta según guía UCI"""
    print("📄 Abriendo documento original...")
    doc = Document("Informe_Tecnico_PID_Gestion_Pedidos.docx")
    
    print("🔍 Identificando secciones actuales...")
    
    # Encontrar índices de las secciones principales
    indices = {
        'introduccion': -1,
        'cap1': -1,
        'cap2': -1,
        'cap3': -1,
        'cap4': -1,
        'conclusiones': -1,
        'recomendaciones': -1,
        'referencias': -1
    }
    
    for i, para in enumerate(doc.paragraphs):
        texto = para.text.strip().upper()
        
        if 'INTRODUCCIÓN' in texto and para.style.name.startswith('Heading'):
            indices['introduccion'] = i
        elif 'CAPÍTULO I' in texto and 'ESTADO DEL ARTE' in texto:
            indices['cap1'] = i
        elif 'CAPÍTULO II' in texto and 'MODELADO' in texto:
            indices['cap2'] = i
        elif 'CAPÍTULO III' in texto and 'DISEÑO' in texto:
            indices['cap3'] = i
        elif 'CAPÍTULO IV' in texto and 'VALIDACIÓN' in texto:
            indices['cap4'] = i
        elif 'CONCLUSIONES' in texto and para.style.name.startswith('Heading') and indices['conclusiones'] == -1:
            indices['conclusiones'] = i
        elif 'RECOMENDACIONES' in texto and para.style.name.startswith('Heading'):
            indices['recomendaciones'] = i
        elif 'REFERENCIAS BIBLIOGRÁFICAS' in texto or 'REFERENCIAS BIBLIOGRAFICAS' in texto:
            indices['referencias'] = i
    
    print("\n📊 Índices encontrados:")
    for key, val in indices.items():
        print(f"   {key}: párrafo {val}")
    
    print("\n🔧 Creando nuevo documento reestructurado...")
    nuevo_doc = Document()
    
    # Copiar estilos del documento original
    nuevo_doc.styles._element = doc.styles._element
    
    print("\n📋 Copiando secciones preliminares...")
    # Copiar todo hasta la introducción (portada, resumen, abstract)
    for i in range(indices['introduccion']):
        para_original = doc.paragraphs[i]
        nuevo_para = nuevo_doc.add_paragraph(para_original.text, style=para_original.style)
        
        # Copiar formato de runs
        nuevo_para.clear()
        for run in para_original.runs:
            nuevo_run = nuevo_para.add_run(run.text)
            if run.bold:
                nuevo_run.bold = True
            if run.italic:
                nuevo_run.italic = True
            if run.font.size:
                nuevo_run.font.size = run.font.size
            if run.font.name:
                nuevo_run.font.name = run.font.name
        
        # Copiar alineación
        nuevo_para.alignment = para_original.alignment
    
    print("📝 Copiando INTRODUCCIÓN...")
    # Copiar introducción completa
    for i in range(indices['introduccion'], indices['cap1']):
        para_original = doc.paragraphs[i]
        nuevo_para = nuevo_doc.add_paragraph(para_original.text, style=para_original.style)
        
        # Copiar formato
        nuevo_para.clear()
        for run in para_original.runs:
            nuevo_run = nuevo_para.add_run(run.text)
            if run.bold:
                nuevo_run.bold = True
            if run.italic:
                nuevo_run.italic = True
            if run.font.size:
                nuevo_run.font.size = run.font.size
            if run.font.name:
                nuevo_run.font.name = run.font.name
        
        nuevo_para.alignment = para_original.alignment
    
    print("\n🏗️ Creando nueva estructura de DESARROLLO...")
    
    # NUEVO CAPÍTULO ÚNICO
    nuevo_doc.add_heading("CAPÍTULO I. DISEÑO E IMPLEMENTACIÓN DEL SISTEMA DE GESTIÓN DE PEDIDOS CON TRAZABILIDAD", level=1)
    
    # Párrafo introductorio del capítulo
    p_intro = nuevo_doc.add_paragraph(
        "Este capítulo presenta el proceso completo de desarrollo del sistema de gestión de pedidos, "
        "desde la fundamentación teórica hasta la validación de la solución implementada. Se estructura "
        "en ocho epígrafes que abordan: los fundamentos conceptuales, el análisis de soluciones existentes, "
        "el diagnóstico de la situación actual, las tecnologías seleccionadas, la propuesta de solución, "
        "los requisitos identificados, el diseño e implementación del sistema, y finalmente la validación "
        "de los resultados obtenidos."
    )
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # EPÍGRAFE I: Contenido del actual Capítulo I hasta la sección 1.2
    print("   ✅ I. Conceptos fundamentales...")
    nuevo_doc.add_heading("I. Conceptos asociados a la gestión de pedidos", level=2)
    
    # Copiar desde inicio Cap I hasta sección 1.2
    for i in range(indices['cap1'] + 1, indices['cap1'] + 40):
        if i >= len(doc.paragraphs):
            break
        para = doc.paragraphs[i]
        if 'Análisis de Soluciones Existentes' in para.text or '1.2' in para.text:
            break
        
        nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
        nuevo_para.alignment = para.alignment
        
        # Ajustar nivel de encabezados
        if para.style.name == 'Heading 2':
            nuevo_para.style = 'Heading 3'
        elif para.style.name == 'Heading 3':
            nuevo_para.style = 'Heading 4'
    
    # EPÍGRAFE II: Soluciones existentes
    print("   ✅ II. Soluciones informáticas existentes...")
    nuevo_doc.add_heading("II. Soluciones informáticas para la gestión de pedidos con trazabilidad", level=2)
    
    # Buscar sección 1.2 y copiar hasta 1.3
    inicio_12 = -1
    fin_12 = -1
    for i in range(indices['cap1'], indices['cap2']):
        if '1.2' in doc.paragraphs[i].text and 'Análisis' in doc.paragraphs[i].text:
            inicio_12 = i
        if '1.3' in doc.paragraphs[i].text or '1.4' in doc.paragraphs[i].text:
            fin_12 = i
            break
    
    if inicio_12 > 0 and fin_12 > 0:
        for i in range(inicio_12 + 1, fin_12):
            para = doc.paragraphs[i]
            nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
            nuevo_para.alignment = para.alignment
            
            if para.style.name == 'Heading 2':
                nuevo_para.style = 'Heading 3'
            elif para.style.name == 'Heading 3':
                nuevo_para.style = 'Heading 4'
    
    # Copiar tabla comparativa si existe
    for tabla in doc.tables[:2]:
        nuevo_doc.add_paragraph()  # Espacio antes
        nueva_tabla = nuevo_doc.add_table(rows=len(tabla.rows), cols=len(tabla.columns))
        nueva_tabla.style = 'Light Grid Accent 1'
        
        for i, fila in enumerate(tabla.rows):
            for j, celda in enumerate(fila.cells):
                nueva_tabla.rows[i].cells[j].text = celda.text
    
    # EPÍGRAFE III: Diagnóstico (contenido del Cap II inicial)
    print("   ✅ III. Proceso de gestión de pedidos en el contexto de estudio...")
    nuevo_doc.add_heading("III. El proceso de gestión de pedidos en pequeñas y medianas empresas cubanas", level=2)
    
    p_diag = nuevo_doc.add_paragraph(
        "Se realizó un análisis del proceso actual de gestión de pedidos en el contexto de pequeñas y medianas "
        "empresas (PYMES) en Cuba, identificando las principales problemáticas y necesidades que justifican "
        "el desarrollo de la solución propuesta. El diagnóstico reveló que la mayoría de estas empresas "
        "utilizan herramientas genéricas como hojas de cálculo y mensajería instantánea, lo que resulta en "
        "errores de inventario, pérdida de trazabilidad de transacciones y dificultades en la generación de "
        "reportes financieros. Estas limitaciones impactan negativamente en la eficiencia operativa y la "
        "toma de decisiones estratégicas."
    )
    p_diag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # EPÍGRAFE IV: Tecnologías (resto del Cap I)
    print("   ✅ IV. Tecnologías seleccionadas...")
    nuevo_doc.add_heading("IV. Tecnologías informáticas para el desarrollo del sistema de gestión de pedidos", level=2)
    
    # Buscar sección 1.3 o 1.4 (tecnologías)
    inicio_tech = -1
    for i in range(indices['cap1'], indices['cap2']):
        if '1.3' in doc.paragraphs[i].text or '1.4' in doc.paragraphs[i].text or 'Tecnología' in doc.paragraphs[i].text:
            inicio_tech = i
            break
    
    if inicio_tech > 0:
        for i in range(inicio_tech, indices['cap2']):
            para = doc.paragraphs[i]
            if para.text.strip():
                nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
                nuevo_para.alignment = para.alignment
                
                if para.style.name == 'Heading 2':
                    nuevo_para.style = 'Heading 3'
                elif para.style.name == 'Heading 3':
                    nuevo_para.style = 'Heading 4'
    
    # EPÍGRAFE V: Descripción de la solución
    print("   ✅ V. Descripción de la solución propuesta...")
    nuevo_doc.add_heading("V. Descripción de la solución informática propuesta", level=2)
    
    p_sol = nuevo_doc.add_paragraph(
        "Se propone el desarrollo de un sistema web basado en arquitectura en capas que integra: "
        "(1) una capa de presentación mediante API RESTful con FastAPI, (2) una capa de lógica de negocio "
        "que implementa las reglas de validación y control de acceso basado en roles (RBAC), (3) una capa "
        "de acceso a datos utilizando SQLAlchemy como ORM, y (4) una capa de persistencia sobre PostgreSQL. "
        "El sistema gestiona integralmente el ciclo de vida de los pedidos: creación, validación de stock, "
        "registro de pagos parciales o totales, actualización automática de inventario, auditoría de acciones "
        "y generación de reportes estadísticos. La autenticación se implementa mediante tokens JWT, garantizando "
        "seguridad en las comunicaciones y control de sesiones."
    )
    p_sol.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # EPÍGRAFE VI: Requisitos (contenido del Cap II)
    print("   ✅ VI. Ingeniería de requisitos...")
    nuevo_doc.add_heading("VI. Requisitos, análisis y modelado del sistema de gestión de pedidos", level=2)
    
    # Copiar todo el capítulo II
    for i in range(indices['cap2'] + 1, indices['cap3']):
        para = doc.paragraphs[i]
        nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
        nuevo_para.alignment = para.alignment
        
        # Ajustar niveles de encabezado
        if para.style.name == 'Heading 1':
            nuevo_para.style = 'Heading 2'
        elif para.style.name == 'Heading 2':
            nuevo_para.style = 'Heading 3'
        elif para.style.name == 'Heading 3':
            nuevo_para.style = 'Heading 4'
    
    # Copiar tablas del Cap II
    print("   📊 Copiando tablas de requisitos...")
    for tabla_idx in range(2, 4):
        if tabla_idx < len(doc.tables):
            tabla = doc.tables[tabla_idx]
            nuevo_doc.add_paragraph()
            nueva_tabla = nuevo_doc.add_table(rows=len(tabla.rows), cols=len(tabla.columns))
            nueva_tabla.style = 'Light Grid Accent 1'
            
            for i, fila in enumerate(tabla.rows):
                for j, celda in enumerate(fila.cells):
                    nueva_tabla.rows[i].cells[j].text = celda.text
    
    # EPÍGRAFE VII: Diseño e implementación (Cap III)
    print("   ✅ VII. Diseño e implementación...")
    nuevo_doc.add_heading("VII. Diseño e implementación del almacenamiento, procesamiento y transmisión de datos", level=2)
    
    # Copiar todo el capítulo III incluyendo la sección 3.5 (si está mal ubicada, la buscamos)
    for i in range(indices['cap3'] + 1, indices['cap4']):
        para = doc.paragraphs[i]
        nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
        nuevo_para.alignment = para.alignment
        
        if para.style.name == 'Heading 1':
            nuevo_para.style = 'Heading 2'
        elif para.style.name == 'Heading 2':
            nuevo_para.style = 'Heading 3'
        elif para.style.name == 'Heading 3':
            nuevo_para.style = 'Heading 4'
    
    # Buscar y copiar sección 3.5 si está después
    print("   🔍 Buscando sección 3.5 de ejemplos de código...")
    inicio_35 = -1
    fin_35 = -1
    for i in range(indices['cap4'], len(doc.paragraphs)):
        if '3.5' in doc.paragraphs[i].text and 'Ejemplos' in doc.paragraphs[i].text:
            inicio_35 = i
        if inicio_35 > 0 and ('CAPÍTULO' in doc.paragraphs[i].text.upper() or 'CONCLUSIONES' in doc.paragraphs[i].text.upper()):
            fin_35 = i
            break
    
    if inicio_35 > 0:
        print(f"   📝 Copiando sección 3.5 desde párrafo {inicio_35}...")
        for i in range(inicio_35, fin_35 if fin_35 > 0 else inicio_35 + 30):
            if i >= len(doc.paragraphs):
                break
            para = doc.paragraphs[i]
            if 'CONCLUSIONES' in para.text.upper() or 'RECOMENDACIONES' in para.text.upper():
                break
            
            nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
            nuevo_para.alignment = para.alignment
            
            if para.style.name == 'Heading 2':
                nuevo_para.style = 'Heading 3'
            elif para.style.name == 'Heading 3':
                nuevo_para.style = 'Heading 4'
    
    # Copiar diagrama ER
    print("   🖼️ Copiando diagrama ER...")
    for i in range(indices['cap3'], indices['cap4']):
        para = doc.paragraphs[i]
        if para._element.xpath('.//pic:pic'):
            nuevo_doc.add_paragraph()
            # Copiar imagen (simplificado - python-docx tiene limitaciones)
            p_fig = nuevo_doc.add_paragraph()
            p_fig.add_run("Figura 3.1: Diagrama entidad-relación del sistema de gestión de pedidos").italic = True
            p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # EPÍGRAFE VIII: Validación (Cap IV)
    print("   ✅ VIII. Verificación y validación...")
    nuevo_doc.add_heading("VIII. Verificación y validación del sistema de gestión de pedidos", level=2)
    
    # Copiar todo el capítulo IV
    for i in range(indices['cap4'] + 1, indices['conclusiones']):
        para = doc.paragraphs[i]
        # Evitar duplicar sección 3.5 si estaba después
        if '3.5' in para.text and 'Ejemplos' in para.text:
            break
        
        nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
        nuevo_para.alignment = para.alignment
        
        if para.style.name == 'Heading 1':
            nuevo_para.style = 'Heading 2'
        elif para.style.name == 'Heading 2':
            nuevo_para.style = 'Heading 3'
        elif para.style.name == 'Heading 3':
            nuevo_para.style = 'Heading 4'
    
    # Copiar tablas del Cap IV
    print("   📊 Copiando tablas de pruebas...")
    for tabla_idx in range(len(doc.tables) - 2, len(doc.tables)):
        if tabla_idx >= 0 and tabla_idx < len(doc.tables):
            tabla = doc.tables[tabla_idx]
            nuevo_doc.add_paragraph()
            nueva_tabla = nuevo_doc.add_table(rows=len(tabla.rows), cols=len(tabla.columns))
            nueva_tabla.style = 'Light Grid Accent 1'
            
            for i, fila in enumerate(tabla.rows):
                for j, celda in enumerate(fila.cells):
                    nueva_tabla.rows[i].cells[j].text = celda.text
    
    print("\n📝 Copiando CONCLUSIONES, RECOMENDACIONES y REFERENCIAS...")
    # Copiar conclusiones, recomendaciones y referencias
    for i in range(indices['conclusiones'], len(doc.paragraphs)):
        para = doc.paragraphs[i]
        nuevo_para = nuevo_doc.add_paragraph(para.text, style=para.style)
        nuevo_para.alignment = para.alignment
    
    print("\n💾 Guardando documento reestructurado...")
    nuevo_doc.save("Informe_Tecnico_PID_Gestion_Pedidos_REESTRUCTURADO.docx")
    
    print("\n" + "="*80)
    print("✅ DOCUMENTO REESTRUCTURADO EXITOSAMENTE")
    print("="*80)
    print("\n📄 Archivo generado: Informe_Tecnico_PID_Gestion_Pedidos_REESTRUCTURADO.docx")
    print("\n📋 Nueva estructura (según guía UCI):")
    print("   ✅ INTRODUCCIÓN")
    print("   ✅ CAPÍTULO I. Diseño e implementación del sistema...")
    print("      I.    Conceptos fundamentales")
    print("      II.   Soluciones informáticas existentes")
    print("      III.  Diagnóstico del contexto actual")
    print("      IV.   Tecnologías seleccionadas")
    print("      V.    Descripción de la solución")
    print("      VI.   Ingeniería de requisitos")
    print("      VII.  Diseño e implementación")
    print("      VIII. Verificación y validación")
    print("   ✅ CONCLUSIONES")
    print("   ✅ RECOMENDACIONES")
    print("   ✅ REFERENCIAS BIBLIOGRÁFICAS")
    
    print("\n⚠️ REVISAR MANUALMENTE:")
    print("   1. Numeración de tablas y figuras")
    print("   2. Referencias cruzadas en el texto")
    print("   3. Formato de imágenes (el diagrama ER debe reinsertarse)")
    print("   4. Generar tabla de contenidos automática en Word")
    print("   5. Ajustar cualquier marcador rojo que haya quedado")


if __name__ == "__main__":
    try:
        crear_documento_reestructurado()
    except Exception as e:
        print(f"\n❌ Error durante la reestructuración: {e}")
        import traceback
        traceback.print_exc()
