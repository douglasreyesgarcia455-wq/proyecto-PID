from docx import Document
import re

# Rutas de archivos
TXT_PATH = 'epigrafe_viii_validacion.txt'
DOCX_PATH = 'Informe_Tecnico_PID_NUEVO_V1.docx'

# Leer bloques de texto del archivo txt (secciones separadas por encabezados ## o #)
def leer_bloques_txt(path):
    with open(path, encoding='utf-8') as f:
        contenido = f.read()
    # Separar por encabezados de sección
    bloques = re.split(r'(?m)^##?\s', contenido)
    # Eliminar caracteres especiales problemáticos
    bloques_limpios = [re.sub(r'[^\w\s.,:;¡!¿?\-\(\)\[\]/@%áéíóúÁÉÍÓÚñÑ"]+', '', b).strip() for b in bloques if b.strip()]
    return bloques_limpios

# Insertar cada bloque como un nuevo párrafo en el documento Word
def insertar_bloques_en_docx(bloques, docx_path):
    doc = Document(docx_path)
    for bloque in bloques:
        doc.add_paragraph(bloque)
        doc.add_paragraph('')  # Espacio entre bloques
    doc.save(docx_path)

if __name__ == '__main__':
    bloques = leer_bloques_txt(TXT_PATH)
    insertar_bloques_en_docx(bloques, DOCX_PATH)
    print('Epígrafe VIII insertado correctamente en el documento Word.')
