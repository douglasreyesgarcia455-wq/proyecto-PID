"""Script para agregar código real en las secciones marcadas como PENDIENTE"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def agregar_codigo_formateado(doc, codigo, titulo, explicacion=""):
    """Agrega un bloque de código con formato"""
    if explicacion:
        p = doc.add_paragraph()
        p.add_run(explicacion).font.size = Pt(12)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Título del código
    p_titulo = doc.add_paragraph()
    run_titulo = p_titulo.add_run(titulo)
    run_titulo.bold = True
    run_titulo.font.size = Pt(11)
    p_titulo.paragraph_format.space_before = Pt(6)
    p_titulo.paragraph_format.space_after = Pt(3)
    
    # Código
    p_codigo = doc.add_paragraph()
    run_codigo = p_codigo.add_run(codigo)
    run_codigo.font.name = 'Consolas'
    run_codigo.font.size = Pt(9)
    run_codigo.font.color.rgb = RGBColor(0, 0, 0)
    
    p_codigo.paragraph_format.left_indent = Inches(0.25)
    p_codigo.paragraph_format.right_indent = Inches(0.25)
    p_codigo.paragraph_format.space_before = Pt(6)
    p_codigo.paragraph_format.space_after = Pt(6)
    p_codigo.paragraph_format.line_spacing = 1.0
    
    return p_codigo


def reemplazar_pendientes():
    """Reemplaza marcadores PENDIENTE con código real"""
    print("📄 Abriendo documento reestructurado...")
    doc = Document("Informe_Tecnico_PID_Gestion_Pedidos_REESTRUCTURADO.docx")
    
    # Códigos a insertar
    codigo_modelo = '''from sqlalchemy import Column, Integer, String, Boolean, DateTime
from datetime import datetime
from src.core.database import Base

class Usuario(Base):
    __tablename__ = "usuarios"
    
    id = Column(Integer, primary_key=True, index=True)
    username = Column(String(50), unique=True, nullable=False, index=True)
    email = Column(String(100), unique=True, nullable=False, index=True)
    hashed_password = Column(String(255), nullable=False)
    rol = Column(String(11), nullable=False)  # admin, vendedor, supervisor
    is_active = Column(Boolean, default=True, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow, nullable=False)
    updated_at = Column(DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)'''
    
    codigo_endpoint = '''from fastapi import APIRouter, Depends, HTTPException, status
from sqlalchemy.orm import Session
from src.core.database import get_db
from src.modules.auth.schema import LoginRequest, TokenResponse
from src.modules.auth.service import AuthService

router = APIRouter(prefix="/api/auth", tags=["auth"])

@router.post("/login", response_model=TokenResponse)
def login(credentials: LoginRequest, db: Session = Depends(get_db)):
    """
    Endpoint de autenticación - retorna token JWT
    
    Args:
        credentials: Objeto con username y password
        db: Sesión de base de datos inyectada
        
    Returns:
        TokenResponse con access_token, user_id, username y rol
        
    Raises:
        HTTPException 401: Credenciales incorrectas
        HTTPException 400: Usuario inactivo
    """
    # Validar credenciales y obtener usuario
    user = AuthService.authenticate_user(db, credentials.username, credentials.password)
    
    # Generar token JWT
    access_token = AuthService.create_user_token(user)
    
    # Retornar respuesta
    return TokenResponse(
        access_token=access_token,
        user_id=user.id,
        username=user.username,
        rol=user.rol
    )'''
    
    codigo_decorador = '''from functools import wraps
from fastapi import HTTPException, status, Depends
from src.core.auth import get_current_user
from src.modules.users.model import Usuario

def require_role(*allowed_roles: str):
    """
    Decorador para validar roles de usuario en endpoints
    
    Args:
        *allowed_roles: Roles permitidos (admin, supervisor, vendedor)
        
    Usage:
        @router.get("/protected")
        @require_role("admin", "supervisor")
        def protected_endpoint(current_user: Usuario = Depends(get_current_user)):
            return {"message": "Acceso permitido"}
    """
    def decorator(func):
        @wraps(func)
        async def wrapper(*args, current_user: Usuario = Depends(get_current_user), **kwargs):
            if current_user.rol not in allowed_roles:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail=f"Acceso denegado. Roles permitidos: {', '.join(allowed_roles)}"
                )
            return await func(*args, current_user=current_user, **kwargs)
        return wrapper
    return decorator'''
    
    codigo_servicio = '''from sqlalchemy.orm import Session
from fastapi import HTTPException
from src.modules.orders.model import Pedido, DetallePedido
from src.modules.products.model import Producto
from src.modules.clients.model import Cliente
from decimal import Decimal

class OrderService:
    
    @staticmethod
    def create_order(db: Session, order_data: OrderCreate) -> Pedido:
        """
        Crea un nuevo pedido con validación de stock y transacción ACID
        
        Args:
            db: Sesión de base de datos
            order_data: Datos del pedido (cliente_id, detalles, pago_inmediato)
            
        Returns:
            Pedido creado con detalles
            
        Raises:
            HTTPException 404: Cliente o producto no encontrado
            HTTPException 400: Stock insuficiente
        """
        # 1. Verificar que el cliente existe
        client = db.query(Cliente).filter(Cliente.id == order_data.cliente_id).first()
        if not client:
            raise HTTPException(status_code=404, detail="Cliente no encontrado")
        
        # 2. Calcular total y validar stock de todos los productos
        total = Decimal(0)
        detalles_validados = []
        
        for detalle in order_data.detalles:
            producto = db.query(Producto).filter(Producto.id == detalle.producto_id).first()
            if not producto:
                raise HTTPException(
                    status_code=404, 
                    detail=f"Producto {detalle.producto_id} no encontrado"
                )
            
            # Validar stock disponible
            if producto.stock < detalle.cantidad:
                raise HTTPException(
                    status_code=400,
                    detail=f"Stock insuficiente para '{producto.nombre}'. "
                           f"Disponible: {producto.stock}, Solicitado: {detalle.cantidad}"
                )
            
            # Calcular subtotal
            precio = detalle.precio_unitario or Decimal(str(producto.precio_venta))
            subtotal = precio * detalle.cantidad
            total += subtotal
            
            detalles_validados.append({
                "producto_id": detalle.producto_id,
                "cantidad": detalle.cantidad,
                "precio_unitario": precio,
                "subtotal": subtotal,
                "producto": producto
            })
        
        # 3. Crear pedido en base de datos
        estado_inicial = "pagado" if order_data.pago_inmediato else "pendiente"
        total_pagado = total if order_data.pago_inmediato else Decimal(0)
        
        db_order = Pedido(
            cliente_id=order_data.cliente_id,
            estado=estado_inicial,
            total=total,
            total_pagado=total_pagado
        )
        db.add(db_order)
        db.flush()  # Obtener ID del pedido sin hacer commit
        
        # 4. Crear detalles del pedido y actualizar inventario
        for detalle in detalles_validados:
            producto = detalle.pop("producto")
            
            # Crear detalle
            db_detalle = DetallePedido(pedido_id=db_order.id, **detalle)
            db.add(db_detalle)
            
            # Reducir stock del producto
            producto.stock -= detalle["cantidad"]
        
        # 5. Commit de la transacción (todo o nada)
        db.commit()
        db.refresh(db_order)
        
        return db_order'''
    
    print("🔍 Buscando marcadores PENDIENTE...")
    
    indices_pendientes = []
    for i, para in enumerate(doc.paragraphs):
        if "[PENDIENTE" in para.text.upper() and "CÓDIGO" in para.text.upper():
            indices_pendientes.append((i, para.text[:100]))
    
    print(f"\n📋 Encontrados {len(indices_pendientes)} marcadores:")
    for idx, texto in indices_pendientes:
        print(f"   - Párrafo {idx}: {texto}...")
    
    if not indices_pendientes:
        print("\n⚠️ No se encontraron marcadores [PENDIENTE en el documento")
        print("   Agregando sección con ejemplos de código al final del epígrafe VII...")
    
    # Buscar el epígrafe VII para agregar los códigos
    indice_epigrafe_vii = -1
    indice_epigrafe_viii = -1
    
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("VII.") and "Diseño" in para.text:
            indice_epigrafe_vii = i
        if para.text.strip().startswith("VIII.") and "Verificación" in para.text:
            indice_epigrafe_viii = i
            break
    
    print(f"\n📍 Epígrafe VII encontrado en párrafo: {indice_epigrafe_vii}")
    print(f"📍 Epígrafe VIII encontrado en párrafo: {indice_epigrafe_viii}")
    
    if indice_epigrafe_vii < 0:
        print("❌ No se encontró el epígrafe VII")
        return
    
    # Como python-docx no permite insertar en posiciones específicas fácilmente,
    # agregamos al final del documento y luego indicamos que hay que moverlo
    
    print("\n✍️ Agregando ejemplos de código...")
    
    # Agregar subsección de ejemplos
    doc.add_page_break()
    doc.add_heading("VII.5. Ejemplos de Implementación", level=3)
    
    p_intro = doc.add_paragraph(
        "A continuación se presentan ejemplos representativos de la implementación del sistema, "
        "mostrando la estructura de las capas de datos, lógica de negocio y presentación, así como "
        "los mecanismos de seguridad implementados."
    )
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Ejemplo 1: Modelo
    doc.add_heading("VII.5.1. Capa de Datos: Modelo Usuario", level=4)
    agregar_codigo_formateado(
        doc,
        codigo_modelo,
        "Código 1: Definición del modelo Usuario con SQLAlchemy ORM",
        "El siguiente código muestra la definición del modelo Usuario utilizando SQLAlchemy ORM. "
        "Se definen columnas con restricciones de integridad (unique, nullable), índices para "
        "optimizar búsquedas por username y email, y campos de auditoría (created_at, updated_at) "
        "que se actualizan automáticamente."
    )
    
    # Ejemplo 2: Endpoint y decorador
    doc.add_heading("VII.5.2. Capa de Presentación: Endpoint de Autenticación", level=4)
    agregar_codigo_formateado(
        doc,
        codigo_endpoint,
        "Código 2: Endpoint POST /api/auth/login con validación de credenciales",
        "Este endpoint maneja la autenticación de usuarios. Recibe credenciales mediante POST, "
        "valida usuario y contraseña a través del servicio de autenticación (que verifica el hash bcrypt), "
        "y retorna un token JWT con información del usuario. El token incluye claims (sub, username, rol) "
        "que se utilizan para autorización en endpoints protegidos."
    )
    
    # Agregar decorador de roles
    doc.add_paragraph()
    agregar_codigo_formateado(
        doc,
        codigo_decorador,
        "Código 3: Decorador require_role para control de acceso basado en roles (RBAC)",
        "El decorador require_role implementa el control de acceso a nivel de endpoint. Se aplica "
        "a funciones que requieren roles específicos, verificando que el usuario autenticado tenga "
        "uno de los roles permitidos. Si el rol no es válido, retorna HTTP 403 Forbidden."
    )
    
    # Ejemplo 3: Servicio
    doc.add_heading("VII.5.3. Capa de Lógica: Servicio de Creación de Pedidos", level=4)
    agregar_codigo_formateado(
        doc,
        codigo_servicio,
        "Código 4: Método create_order() con validación de stock y transacciones ACID",
        "Este método implementa la lógica completa de creación de pedidos: (1) valida existencia del cliente, "
        "(2) verifica stock disponible de todos los productos solicitados, (3) calcula totales con precios "
        "unitarios, (4) crea el pedido y sus detalles en la base de datos, (5) actualiza inventario de productos. "
        "Todo se ejecuta dentro de una transacción única (db.commit() al final), garantizando propiedades ACID: "
        "si cualquier paso falla, toda la operación se revierte automáticamente (rollback)."
    )
    
    # Párrafo de cierre
    p_conclusion = doc.add_paragraph(
        "Estos ejemplos demuestran la separación clara de responsabilidades implementada en el sistema: "
        "los modelos definen la estructura de datos y mapeo objeto-relacional, los servicios contienen "
        "la lógica de negocio con validaciones y manejo de transacciones, y los endpoints exponen "
        "funcionalidad a través de una API RESTful con control de acceso. Esta arquitectura facilita "
        "el mantenimiento, pruebas unitarias y escalabilidad del sistema, permitiendo modificar cada "
        "capa de forma independiente sin afectar las demás."
    )
    p_conclusion.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_conclusion.paragraph_format.space_before = Pt(12)
    
    print("\n💾 Guardando documento con código agregado...")
    doc.save("Informe_Tecnico_PID_Gestion_Pedidos_REESTRUCTURADO.docx")
    
    print("\n" + "="*80)
    print("✅ CÓDIGO AGREGADO EXITOSAMENTE")
    print("="*80)
    print("\n📋 Ejemplos agregados:")
    print("   ✅ Código 1: Modelo Usuario (SQLAlchemy)")
    print("   ✅ Código 2: Endpoint /api/auth/login (FastAPI)")
    print("   ✅ Código 3: Decorador @require_role (RBAC)")
    print("   ✅ Código 4: Servicio crear_pedido() (Transacciones ACID)")
    
    print("\n📍 UBICACIÓN:")
    print("   Los ejemplos se agregaron al final del documento")
    print("   ⚠️ IMPORTANTE: Mover la sección 'VII.5. Ejemplos de Implementación'")
    print("   para que quede ANTES del epígrafe VIII (Verificación y validación)")
    
    print("\n💡 Pasos manuales:")
    print("   1. Abre el documento")
    print("   2. Busca 'VII.5. Ejemplos de Implementación' (al final)")
    print("   3. Selecciona toda la sección VII.5 (incluye los 4 códigos)")
    print("   4. Corta (Ctrl+X)")
    print("   5. Busca 'VIII. Verificación y validación'")
    print("   6. Posiciona cursor ANTES de ese título")
    print("   7. Pega (Ctrl+V)")
    print("   8. Guarda")


if __name__ == "__main__":
    try:
        reemplazar_pendientes()
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
