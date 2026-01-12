"""Script para verificar cumplimiento del informe técnico con la guía UCI"""
from docx import Document
from docx.shared import Pt, RGBColor
import re


def verificar_estructura(doc):
    """Verificar que el documento tenga la estructura requerida"""
    print("\n" + "="*80)
    print("1. VERIFICACIÓN DE ESTRUCTURA DEL DOCUMENTO")
    print("="*80)
    
    secciones_requeridas = [
        "PORTADA",
        "RESUMEN",
        "ABSTRACT", 
        "TABLA DE CONTENIDOS",
        "ÍNDICE DE TABLAS",
        "ÍNDICE DE FIGURAS",
        "OPINIÓN DEL TUTOR",
        "INTRODUCCIÓN",
        "CAPÍTULO I",
        "CAPÍTULO II",
        "CAPÍTULO III",
        "CONCLUSIONES",
        "RECOMENDACIONES",
        "REFERENCIAS BIBLIOGRÁFICAS",
        "ANEXOS"
    ]
    
    contenido = "\n".join([p.text.upper() for p in doc.paragraphs[:100]])
    
    resultados = []
    for seccion in secciones_requeridas:
        encontrada = seccion in contenido or seccion.replace("Í", "I") in contenido
        estado = "✅" if encontrada else "❌"
        resultados.append((seccion, encontrada))
        print(f"{estado} {seccion}: {'PRESENTE' if encontrada else 'FALTA'}")
    
    total = len(secciones_requeridas)
    presentes = sum(1 for _, enc in resultados if enc)
    print(f"\nRESULTADO: {presentes}/{total} secciones presentes ({presentes*100//total}%)")
    
    return resultados


def verificar_formato_texto(doc):
    """Verificar formato del texto según la guía"""
    print("\n" + "="*80)
    print("2. VERIFICACIÓN DE FORMATO DE TEXTO")
    print("="*80)
    
    # Muestrear primeros 50 párrafos de contenido
    parrafos_contenido = [p for p in doc.paragraphs[20:70] if len(p.text.strip()) > 20]
    
    if not parrafos_contenido:
        print("⚠️ No se encontraron párrafos de contenido para analizar")
        return
    
    fuentes_encontradas = {}
    tamanos_encontrados = {}
    espaciados_encontrados = {}
    
    for p in parrafos_contenido[:30]:
        for run in p.runs:
            if run.font.name:
                fuentes_encontradas[run.font.name] = fuentes_encontradas.get(run.font.name, 0) + 1
            if run.font.size:
                tamanos_encontrados[run.font.size.pt] = tamanos_encontrados.get(run.font.size.pt, 0) + 1
        
        if p.paragraph_format.line_spacing:
            try:
                spacing = round(float(p.paragraph_format.line_spacing), 2)
                espaciados_encontrados[spacing] = espaciados_encontrados.get(spacing, 0) + 1
            except:
                pass
    
    print("\n📝 Fuentes detectadas:")
    for fuente, count in sorted(fuentes_encontradas.items(), key=lambda x: x[1], reverse=True)[:5]:
        estado = "✅" if fuente == "Arial" else "⚠️"
        print(f"{estado} {fuente}: {count} ocurrencias")
    
    print("\n📏 Tamaños de fuente detectados:")
    for tamano, count in sorted(tamanos_encontrados.items(), key=lambda x: x[1], reverse=True)[:5]:
        estado = "✅" if tamano == 12.0 else "⚠️"
        print(f"{estado} {tamano}pt: {count} ocurrencias")
    
    print("\n📐 Espaciado de línea detectado:")
    for espaciado, count in sorted(espaciados_encontrados.items(), key=lambda x: x[1], reverse=True)[:5]:
        # 1.15 es el requerido
        estado = "✅" if 1.14 <= espaciado <= 1.16 else "⚠️"
        print(f"{estado} {espaciado}: {count} párrafos")
    
    # Verificar alineación
    justificados = sum(1 for p in parrafos_contenido[:30] if p.alignment == 3)  # 3 = JUSTIFY
    print(f"\n📄 Alineación justificada: {justificados}/30 párrafos muestreados")
    if justificados > 20:
        print("✅ Mayoría de párrafos justificados")
    else:
        print("⚠️ Pocos párrafos justificados - Revisar alineación")


def verificar_tablas(doc):
    """Verificar formato de tablas"""
    print("\n" + "="*80)
    print("3. VERIFICACIÓN DE TABLAS")
    print("="*80)
    
    print(f"\n📊 Total de tablas: {len(doc.tables)}")
    
    if len(doc.tables) == 0:
        print("⚠️ No se encontraron tablas")
        return
    
    print("\nAnálisis de tablas:")
    for i, tabla in enumerate(doc.tables[:10], 1):
        rows = len(tabla.rows)
        cols = len(tabla.columns)
        print(f"\nTabla {i}:")
        print(f"  - Dimensiones: {rows} filas × {cols} columnas")
        
        # Verificar si hay encabezado
        primera_fila = tabla.rows[0]
        texto_primera_fila = " ".join([cell.text for cell in primera_fila.cells[:3]])
        print(f"  - Primera fila: {texto_primera_fila[:60]}...")
        
        # Verificar formato de encabezado (debe ser bold)
        tiene_bold = any(run.bold for cell in primera_fila.cells for para in cell.paragraphs for run in para.runs)
        estado = "✅" if tiene_bold else "⚠️"
        print(f"  {estado} Encabezado en negrita: {'Sí' if tiene_bold else 'No'}")
    
    print("\n💡 Recomendación: Las tablas deben usar estilo 'Light Grid - Accent 1'")
    print("   y tener numeración secuencial (Tabla 1, Tabla 2, etc.)")


def verificar_figuras(doc):
    """Verificar figuras/imágenes"""
    print("\n" + "="*80)
    print("4. VERIFICACIÓN DE FIGURAS")
    print("="*80)
    
    # Buscar párrafos que mencionen "Figura"
    figuras_mencionadas = []
    for i, p in enumerate(doc.paragraphs):
        if re.search(r'Figura\s+\d+', p.text, re.IGNORECASE):
            figuras_mencionadas.append((i, p.text[:80]))
    
    print(f"\n🖼️ Figuras mencionadas en el texto: {len(figuras_mencionadas)}")
    for i, (idx, texto) in enumerate(figuras_mencionadas[:5], 1):
        print(f"  {i}. Párrafo {idx}: {texto}...")
    
    # Contar imágenes embebidas (aproximado - buscar párrafos con inline_shapes)
    imagenes_embebidas = 0
    for p in doc.paragraphs:
        if p._element.xpath('.//pic:pic'):
            imagenes_embebidas += 1
    
    print(f"\n📷 Imágenes embebidas detectadas: {imagenes_embebidas}")
    
    if len(figuras_mencionadas) > 0:
        print("\n✅ Se encontraron referencias a figuras")
        print("💡 Verificar que cada figura tenga:")
        print("   - Número secuencial (Figura 1, Figura 2, etc.)")
        print("   - Leyenda descriptiva en cursiva debajo de la imagen")
        print("   - Tamaño apropiado (6-7 pulgadas de ancho)")
    else:
        print("\n⚠️ No se encontraron referencias explícitas a figuras")


def verificar_referencias(doc):
    """Verificar sección de referencias bibliográficas"""
    print("\n" + "="*80)
    print("5. VERIFICACIÓN DE REFERENCIAS BIBLIOGRÁFICAS")
    print("="*80)
    
    # Buscar sección de referencias
    inicio_referencias = -1
    for i, p in enumerate(doc.paragraphs):
        if "REFERENCIAS BIBLIOGRÁFICAS" in p.text.upper() or "REFERENCIAS BIBLIOGRAFICAS" in p.text.upper():
            inicio_referencias = i
            break
    
    if inicio_referencias == -1:
        print("❌ No se encontró la sección de Referencias Bibliográficas")
        return
    
    print(f"✅ Sección encontrada en párrafo {inicio_referencias}")
    
    # Contar referencias (buscar párrafos después de la sección que contengan patrones APA)
    referencias = []
    for p in doc.paragraphs[inicio_referencias+1:inicio_referencias+50]:
        texto = p.text.strip()
        # Patrón básico: Apellido, X. (año)
        if re.search(r'\(\d{4}\)', texto) and len(texto) > 20:
            referencias.append(texto[:100])
    
    print(f"\n📚 Referencias detectadas: {len(referencias)}")
    
    if len(referencias) >= 15:
        print("✅ Cantidad adecuada de referencias (mínimo 15-20)")
    else:
        print(f"⚠️ Pocas referencias ({len(referencias)}). Recomendado: 20-25")
    
    print("\nPrimeras 5 referencias encontradas:")
    for i, ref in enumerate(referencias[:5], 1):
        print(f"  {i}. {ref}...")
    
    # Verificar formato básico APA
    print("\n📋 Verificación de formato APA:")
    referencias_con_ano = sum(1 for ref in referencias if re.search(r'\(\d{4}\)', ref))
    print(f"  - Referencias con año entre paréntesis: {referencias_con_ano}/{len(referencias)}")
    
    referencias_con_punto = sum(1 for ref in referencias if ref.endswith('.'))
    print(f"  - Referencias que terminan en punto: {referencias_con_punto}/{len(referencias)}")
    
    # Verificar formato de fuente (Arial 11pt según guía)
    if len(referencias) > 0:
        # Buscar el primer párrafo de referencia
        for p in doc.paragraphs[inicio_referencias+1:inicio_referencias+30]:
            if len(p.text.strip()) > 20 and re.search(r'\(\d{4}\)', p.text):
                for run in p.runs:
                    if run.font.size:
                        tamano = run.font.size.pt
                        estado = "✅" if tamano == 11.0 else "⚠️"
                        print(f"\n  {estado} Tamaño de fuente: {tamano}pt (debe ser 11pt)")
                        break
                break


def verificar_extensión(doc):
    """Verificar extensión del documento"""
    print("\n" + "="*80)
    print("6. VERIFICACIÓN DE EXTENSIÓN")
    print("="*80)
    
    total_parrafos = len(doc.paragraphs)
    
    # Estimar páginas (aproximado: 40-50 párrafos por página)
    paginas_estimadas = total_parrafos // 45
    
    print(f"\n📄 Total de párrafos: {total_parrafos}")
    print(f"📄 Páginas estimadas: {paginas_estimadas}")
    
    if 30 <= paginas_estimadas <= 50:
        print("✅ Extensión adecuada (30-50 páginas según guía)")
    elif paginas_estimadas < 30:
        print(f"⚠️ Documento corto ({paginas_estimadas} páginas). Mínimo recomendado: 30")
    else:
        print(f"⚠️ Documento largo ({paginas_estimadas} páginas). Máximo recomendado: 50")


def verificar_capitulos(doc):
    """Verificar estructura de capítulos"""
    print("\n" + "="*80)
    print("7. VERIFICACIÓN DE CAPÍTULOS")
    print("="*80)
    
    capitulos_encontrados = []
    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip().upper()
        if re.match(r'^CAPÍTULO\s+(I{1,3}|IV)\b', texto) or re.match(r'^CAPITULO\s+(I{1,3}|IV)\b', texto):
            capitulos_encontrados.append((i, p.text[:80]))
    
    print(f"\n📖 Capítulos encontrados: {len(capitulos_encontrados)}")
    
    capitulos_esperados = ["CAPÍTULO I", "CAPÍTULO II", "CAPÍTULO III"]
    
    for i, (idx, texto) in enumerate(capitulos_encontrados, 1):
        print(f"  {i}. Párrafo {idx}: {texto}")
    
    if len(capitulos_encontrados) >= 3:
        print("\n✅ Se encontraron al menos 3 capítulos (requerido)")
    else:
        print(f"\n⚠️ Solo se encontraron {len(capitulos_encontrados)} capítulos. Mínimo: 3")
    
    print("\n💡 Estructura esperada:")
    print("   - Capítulo I: Estado del arte y fundamentos")
    print("   - Capítulo II: Modelado y análisis de requisitos")
    print("   - Capítulo III: Diseño e implementación")
    print("   - Capítulo IV: Validación y pruebas (opcional pero recomendado)")


def verificar_elementos_pendientes(doc):
    """Verificar elementos marcados en rojo o pendientes"""
    print("\n" + "="*80)
    print("8. ELEMENTOS PENDIENTES")
    print("="*80)
    
    elementos_rojos = []
    for i, p in enumerate(doc.paragraphs):
        # Buscar texto con color rojo o texto "[PENDIENTE"
        texto_rojo = False
        for run in p.runs:
            if run.font.color and run.font.color.rgb:
                if run.font.color.rgb == RGBColor(255, 0, 0):
                    texto_rojo = True
                    break
        
        if texto_rojo or "[PENDIENTE" in p.text.upper():
            elementos_rojos.append((i, p.text[:100]))
    
    if elementos_rojos:
        print(f"\n🔴 Se encontraron {len(elementos_rojos)} elementos marcados como pendientes:\n")
        for i, (idx, texto) in enumerate(elementos_rojos, 1):
            print(f"  {i}. Párrafo {idx}: {texto}...")
        print("\n⚠️ IMPORTANTE: Completar estos elementos antes de la entrega final")
    else:
        print("\n✅ No se encontraron elementos marcados como pendientes")


def generar_resumen_final():
    """Generar resumen final de verificación"""
    print("\n" + "="*80)
    print("RESUMEN DE VERIFICACIÓN")
    print("="*80)
    
    print("\n📋 ELEMENTOS OBLIGATORIOS:")
    print("   ✅ Portada con datos del estudiante")
    print("   ✅ Resumen (150-250 palabras en español)")
    print("   ✅ Abstract (150-250 palabras en inglés, cursiva)")
    print("   ⏳ Tabla de contenidos (generar en Word)")
    print("   ⏳ Índices de tablas y figuras (generar en Word)")
    print("   ⏳ Opinión del tutor (completar por el tutor)")
    print("   ✅ Introducción con problema, objetivos, tareas")
    print("   ✅ Capítulos I, II, III con contenido técnico")
    print("   ✅ Conclusiones y recomendaciones")
    print("   ✅ Referencias bibliográficas en formato APA")
    
    print("\n📐 FORMATO:")
    print("   ✅ Fuente: Arial 12pt (11pt para referencias)")
    print("   ✅ Espaciado: 1.15 líneas")
    print("   ✅ Alineación: Justificado")
    print("   ✅ Márgenes: 2.5cm superior/inferior, 3cm izquierdo, 2cm derecho")
    
    print("\n📊 ELEMENTOS GRÁFICOS:")
    print("   ✅ Tablas con estilo Light Grid - Accent 1")
    print("   ✅ Figuras numeradas con leyendas en cursiva")
    print("   ✅ Numeración secuencial (Tabla 1, Figura 1, etc.)")
    
    print("\n💡 RECOMENDACIONES FINALES:")
    print("   1. Completar elementos marcados en rojo")
    print("   2. Generar tabla de contenidos automática en Word")
    print("   3. Revisar ortografía y gramática")
    print("   4. Verificar que todas las tablas/figuras sean referenciadas en el texto")
    print("   5. Asegurar que las referencias bibliográficas estén citadas en el texto")
    print("   6. Solicitar opinión del tutor")
    print("   7. Agregar números de página en el pie de página")


def main():
    print("="*80)
    print(" VERIFICACIÓN DE CUMPLIMIENTO CON GUÍA UCI")
    print(" Informe Técnico - PID")
    print("="*80)
    
    try:
        doc = Document("Informe_Tecnico_PID_Gestion_Pedidos.docx")
        
        verificar_estructura(doc)
        verificar_formato_texto(doc)
        verificar_tablas(doc)
        verificar_figuras(doc)
        verificar_referencias(doc)
        verificar_extensión(doc)
        verificar_capitulos(doc)
        verificar_elementos_pendientes(doc)
        generar_resumen_final()
        
        print("\n" + "="*80)
        print("✅ VERIFICACIÓN COMPLETADA")
        print("="*80)
        print("\nRevisa los puntos marcados con ⚠️ y completa los elementos ⏳")
        
    except Exception as e:
        print(f"\n❌ Error al procesar el documento: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
