# -*- coding: utf-8 -*-
"""
Script para agregar EPÍGRAFE VII al documento del informe técnico.
Preserva todo el contenido anterior y agrega el nuevo contenido al final.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Nombre del archivo
FILENAME = "Informe_Tecnico_PID_NUEVO_V1.docx"

def agregar_titulo(doc, texto, nivel=1):
    """Agrega un título al documento"""
    p = doc.add_heading(texto, level=nivel)
    return p

def agregar_parrafo(doc, texto):
    """Agrega un párrafo con formato estándar"""
    p = doc.add_paragraph(texto)
    p.style = 'Normal'
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def agregar_imagen(doc, ruta, ancho_inches=6.0):
    """Agrega una imagen al documento"""
    if os.path.exists(ruta):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(ruta, width=Inches(ancho_inches))
        return p
    else:
        print(f"⚠️ Advertencia: No se encontró la imagen: {ruta}")
        return None

def agregar_codigo(doc, codigo, lenguaje="SQL"):
    """Agrega un bloque de código con formato"""
    p = doc.add_paragraph(codigo)
    p.style = 'Normal'
    for run in p.runs:
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def agregar_tabla(doc, headers, rows):
    """Agrega una tabla al documento"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data)
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    return table

def epigrafe_vii(doc):
    """Genera el contenido del Epígrafe VII"""
    
    # TÍTULO PRINCIPAL
    agregar_titulo(doc, "EPÍGRAFE VII: Diseño e implementación del sistema de gestión de pedidos", nivel=1)
    
    agregar_parrafo(doc, 
        "El presente epígrafe detalla el diseño de la arquitectura de datos y la implementación del sistema web "
        "de gestión de pedidos. Se expone el modelo entidad-relación de la base de datos, las tablas con sus "
        "atributos y restricciones, la lógica de negocio implementada y los componentes principales del sistema "
        "desarrollado. Este diseño garantiza la integridad referencial, la trazabilidad de operaciones y el "
        "cumplimiento de los requisitos funcionales y no funcionales establecidos.")
    
    # =============================
    # 7.1 DISEÑO DEL MODELO DE DATOS
    # =============================
    agregar_titulo(doc, "7.1. Diseño del modelo de datos", nivel=2)
    
    agregar_parrafo(doc,
        "El modelo de datos del sistema se diseñó siguiendo el enfoque relacional normalizado hasta la tercera forma "
        "normal (3FN), garantizando la eliminación de redundancias y asegurando la integridad referencial mediante "
        "el uso de claves primarias y foráneas. El diseño contempla 9 entidades principales que representan los "
        "conceptos del dominio: usuarios internos, clientes, contactos de clientes, productos, pedidos con sus detalles, "
        "pagos, devoluciones y un registro de auditoría (logs de acciones). La figura 7.1 muestra el diagrama "
        "entidad-relación que representa las entidades, atributos y relaciones del sistema.")
    
    # INSERTAR DIAGRAMA ER
    agregar_parrafo(doc, "")
    p_figura = doc.add_paragraph()
    p_figura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_fig = p_figura.add_run()
    run_fig.font.bold = True
    run_fig.font.size = Pt(11)
    run_fig.text = "Figura 7.1. Diagrama Entidad-Relación del sistema de gestión de pedidos"
    
    # Insertar imagen del diagrama ER
    agregar_imagen(doc, "diagramas_analisis/diagrama_er_completo.png", ancho_inches=6.5)
    
    agregar_parrafo(doc, "Fuente: Elaboración propia")
    agregar_parrafo(doc, "")
    
    agregar_parrafo(doc,
        "El diagrama entidad-relación presentado utiliza la notación de Chen, donde los rectángulos representan "
        "entidades, los óvalos representan atributos (con subrayado para claves primarias, óvalo discontinuo para "
        "atributos derivados y óvalo doble para multivaluados), y los rombos representan relaciones entre entidades. "
        "Las cardinalidades (1:N, 1:1, N:M) indican la multiplicidad de las asociaciones. El modelo contempla "
        "relaciones de composición (como pedidos que contienen detalles), asociación (clientes que realizan pedidos), "
        "y herencia comportamental (usuarios con diferentes roles).")
    
    # 7.1.1 Descripción de las entidades
    agregar_titulo(doc, "7.1.1. Descripción de las entidades y tablas", nivel=3)
    
    agregar_parrafo(doc,
        "A continuación se describen las entidades principales del modelo de datos y sus correspondientes tablas "
        "en el gestor de base de datos PostgreSQL 16. Cada tabla incluye restricciones de integridad, claves primarias "
        "y foráneas, índices para optimizar consultas frecuentes y valores por defecto según las reglas del negocio.")
    
    # TABLA 7.1: Entidades del sistema
    agregar_parrafo(doc, "")
    p_tabla = doc.add_paragraph()
    p_tabla.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabla = p_tabla.add_run()
    run_tabla.font.bold = True
    run_tabla.font.size = Pt(11)
    run_tabla.text = "Tabla 7.1. Entidades principales del sistema"
    # Tabla pedidos
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE pedidos")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_pedidos = """CREATE TABLE pedidos (
    
    agregar_parrafo(doc,
        "El esquema de la base de datos implementado en PostgreSQL 16 se presenta a continuación. Se utilizan tipos "
        "de datos optimizados para cada campo (SERIAL para PKs autoincrementales, NUMERIC(12,2) para valores monetarios, "
        "TIMESTAMP para fechas y horas, JSONB para datos semiestructurados). Las restricciones ON DELETE CASCADE "
        "garantizan la integridad referencial al eliminar registros padres.")
    
    # Tabla usuarios
    agregar_parrafo(doc, "")
    agregar_codigo(doc, codigo_pedidos)
    agregar_parrafo(doc, "• pedidos: Almacena los pedidos realizados por los clientes, con referencia al usuario que los registró. Permite pagos parciales y controla el estado del pedido.")

    # Tabla detalles_pedido
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    agregar_parrafo(doc,
        "El esquema de la base de datos implementado en PostgreSQL 16 se presenta a continuación. Se utilizan tipos "
        "de datos optimizados para cada campo (SERIAL para PKs autoincrementales, NUMERIC(12,2) para valores monetarios, "
        "TIMESTAMP para fechas y horas, JSONB para datos semiestructurados). Las restricciones ON DELETE CASCADE "
        "garantizan la integridad referencial al eliminar registros padres.")

    # Tabla pedidos
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE pedidos")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_pedidos = (
        "CREATE TABLE pedidos (\n"
        "    id SERIAL PRIMARY KEY,\n"
        "    cliente_id INTEGER NOT NULL REFERENCES clientes(id),\n"
        "    usuario_id INTEGER REFERENCES usuarios(id),\n"
        "    estado VARCHAR(20) DEFAULT 'pendiente',  -- pendiente, pagado, cancelado, devuelto\n"
        "    total NUMERIC(12,2) NOT NULL,\n"
        "    total_pagado NUMERIC(12,2) DEFAULT 0,\n"
        "    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n"
        "    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ");"
    )
    agregar_codigo(doc, codigo_pedidos)
    agregar_parrafo(doc, "• pedidos: Almacena los pedidos realizados por los clientes, con referencia al usuario que los registró. Permite pagos parciales y controla el estado del pedido.")

    # Tabla detalles_pedido
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE detalles_pedido")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_detalles = (
        "CREATE TABLE detalles_pedido (\n"
        "    id SERIAL PRIMARY KEY,\n"
        "    pedido_id INTEGER NOT NULL REFERENCES pedidos(id) ON DELETE CASCADE,\n"
        "    producto_id INTEGER NOT NULL REFERENCES productos(id),\n"
        "    cantidad INTEGER NOT NULL,\n"
        "    precio_unitario NUMERIC(12,2) NOT NULL,\n"
        "    subtotal NUMERIC(12,2) NOT NULL,\n"
        "    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ");"
    )
    agregar_codigo(doc, codigo_detalles)
    agregar_parrafo(doc, "• detalles_pedido: Representa cada línea de producto en un pedido, con cantidad, precio y subtotal. Implementa la relación N:M entre pedidos y productos.")

    # Tabla pagos
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE pagos")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_pagos = (
        "CREATE TABLE pagos (\n"
        "    id SERIAL PRIMARY KEY,\n"
        "    pedido_id INTEGER NOT NULL REFERENCES pedidos(id),\n"
        "    monto NUMERIC(12,2) NOT NULL,\n"
        "    metodo_pago VARCHAR(50),  -- efectivo, transferencia, tarjeta\n"
        "    referencia VARCHAR(255),\n"
        "    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ");"
    )
    agregar_codigo(doc, codigo_pagos)
    agregar_parrafo(doc, "• pagos: Registra los pagos realizados para cada pedido, permitiendo múltiples pagos parciales y diferentes métodos de pago.")

    # Tabla devoluciones
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE devoluciones")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_devoluciones = (
        "CREATE TABLE devoluciones (\n"
        "    id SERIAL PRIMARY KEY,\n"
        "    pedido_id INTEGER UNIQUE NOT NULL REFERENCES pedidos(id),\n"
        "    usuario_id INTEGER REFERENCES usuarios(id),\n"
        "    motivo VARCHAR(255),\n"
        "    descripcion TEXT,\n"
        "    fecha_devolucion TIMESTAMP DEFAULT CURRENT_TIMESTAMP,\n"
        "    productos_devueltos JSONB,  -- [{producto_id, cantidad, precio}]\n"
        "    monto_total NUMERIC(12,2)\n"
        ");"
    )
    agregar_codigo(doc, codigo_devoluciones)
    agregar_parrafo(doc, "• devoluciones: Permite registrar devoluciones de pedidos, con motivo, descripción y productos devueltos en formato JSONB.")

    # Tabla logs_acciones
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE logs_acciones")
    run.font.bold = True
    run.font.size = Pt(11)
    codigo_logs = (
        "CREATE TABLE logs_acciones (\n"
        "    id SERIAL PRIMARY KEY,\n"
        "    usuario_id INTEGER REFERENCES usuarios(id),\n"
        "    endpoint VARCHAR(255),\n"
        "    metodo_http VARCHAR(10),\n"
        "    payload JSONB,\n"
        "    ip_address VARCHAR(45),\n"
        "    user_agent TEXT,\n"
        "    status_code INTEGER,\n"
        "    response_time_ms INTEGER,\n"
        "    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP\n"
        ");"
    )
    agregar_codigo(doc, codigo_logs)
    agregar_parrafo(doc, "• logs_acciones: Tabla de auditoría que registra todas las acciones de los usuarios, incluyendo endpoint, método, payload, IP y tiempos de respuesta.")
    codigo_funcion = """CREATE OR REPLACE FUNCTION calcular_monto_pendiente(pedido_id_param INTEGER)
    agregar_codigo(doc, codigo_clientes)
    
    # Tabla contactos_clientes
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE contactos_clientes")
    run.font.bold = True
    run.font.size = Pt(11)
    
    codigo_contactos = """CREATE TABLE contactos_clientes (
    id SERIAL PRIMARY KEY,
    cliente_id INTEGER NOT NULL REFERENCES clientes(id) ON DELETE CASCADE,
    agregar_codigo(doc, codigo_funcion)
    tipo VARCHAR(20) NOT NULL,  -- telefono, email, whatsapp
    valor VARCHAR(255) NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);"""
    agregar_codigo(doc, codigo_contactos)
    
    # Tabla productos
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE productos")
    run.font.bold = True
    run.font.size = Pt(11)
    
    codigo_productos = """CREATE TABLE productos (
    id SERIAL PRIMARY KEY,
    nombre VARCHAR(255) NOT NULL,
    descripcion TEXT,
    precio_venta NUMERIC(12,2) NOT NULL,
    cantidad INTEGER NOT NULL DEFAULT 0,  -- stock
    stock_minimo INTEGER DEFAULT 5,
    is_active BOOLEAN DEFAULT TRUE,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);"""
    agregar_codigo(doc, codigo_productos)
    
    # Tabla pedidos
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("CREATE TABLE pedidos")
    run.font.bold = True
    run.font.size = Pt(11)
    
    codigo_pedidos = """CREATE TABLE pedidos (
    id SERIAL PRIMARY KEY,
    cliente_id INTEGER NOT NULL REFERENCES clientes(id),
    usuario_id INTEGER REFERENCES usuarios(id),
    estado VARCHAR(20) DEFAULT 'pendiente',  -- pendiente, pagado, cancelado, devuelto
    total NUMERIC(12,2) NOT NULL,
    total_pagado NUMERIC(12,2) DEFAULT 0,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);"""
    agregar_codigo(doc, codigo_pedidos)
    
    # Función SQL calcular_monto_pendiente
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("Función SQL: calcular_monto_pendiente()")
    run.font.bold = True
    run.font.size = Pt(11)
    
    agregar_parrafo(doc,
        "La función calcular_monto_pendiente() se implementó como función almacenada en PostgreSQL para "
        "garantizar cálculos precisos y evitar inconsistencias por redondeo. Retorna el monto pendiente "
        "de un pedido (total - total_pagado) y es utilizada tanto por la API como por reportes.")
    
    codigo_funcion = """CREATE OR REPLACE FUNCTION calcular_monto_pendiente(pedido_id_param INTEGER)
RETURNS NUMERIC(12,2) AS $$
DECLARE
    total_pedido NUMERIC(12,2);
    total_pagado_pedido NUMERIC(12,2);
BEGIN
    SELECT total, total_pagado 
    INTO total_pedido, total_pagado_pedido
    FROM pedidos 
    WHERE id = pedido_id_param;
    
    RETURN total_pedido - total_pagado_pedido;
END;
$$ LANGUAGE plpgsql;"""
    agregar_codigo(doc, codigo_funcion)
    
    # =============================
    # 7.2 IMPLEMENTACIÓN DE REGLAS DE NEGOCIO
    # =============================
    agregar_titulo(doc, "7.2. Implementación de la lógica de negocio", nivel=2)
    
    agregar_parrafo(doc,
        "La lógica de negocio del sistema se implementó en la capa de servicios (service layer) siguiendo el "
        "patrón Repository y el principio de separación de responsabilidades. Cada módulo (usuarios, clientes, "
        "productos, pedidos, pagos, devoluciones) cuenta con un servicio dedicado que encapsula las reglas de "
        "negocio identificadas en el epígrafe III. A continuación se detallan las 9 reglas de negocio principales "
        "y su implementación en código Python utilizando SQLAlchemy ORM y FastAPI.")
    
    # Tabla 7.2: Reglas de negocio
    agregar_parrafo(doc, "")
    p_tabla2 = doc.add_paragraph()
    p_tabla2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabla2 = p_tabla2.add_run()
    run_tabla2.font.bold = True
    run_tabla2.font.size = Pt(11)
    run_tabla2.text = "Tabla 7.2. Reglas de negocio implementadas en el sistema"
    
    tabla_rn = [
        ["RN-01", "Validación de stock", "Verificar stock suficiente antes de crear pedido"],
        ["RN-02", "Reducción automática de stock", "Descontar stock al confirmar pedido"],
        ["RN-03", "Pagos acumulativos", "Permitir múltiples pagos parciales por pedido"],
        ["RN-04", "Cambio automático a estado pagado", "Cambiar estado cuando total_pagado >= total"],
        ["RN-05", "Cálculo exacto de monto pendiente", "Usar función SQL para evitar redondeos"],
        ["RN-06", "Restricción de sobrepago", "Rechazar pagos que excedan el monto pendiente"],
        ["RN-07", "Devolución cambia estado", "Cambiar estado a devuelto al registrar devolución"],
        ["RN-08", "Restauración de inventario", "Sumar productos devueltos al stock"],
        ["RN-09", "Eliminación de pagos en devolución", "Eliminar pagos y resetear total_pagado"]
    ]
    agregar_tabla(doc, 
                  ["Código", "Regla de Negocio", "Descripción"],
                  tabla_rn)
    
    agregar_parrafo(doc, "Fuente: Elaboración propia")
    agregar_parrafo(doc, "")
    
    # Implementación RN-01 y RN-02
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("Implementación de RN-01 y RN-02 (Validación y reducción de stock)")
    run.font.bold = True
    run.font.size = Pt(11)
    
    agregar_parrafo(doc,
        "La validación de stock (RN-01) se realiza antes de confirmar el pedido, verificando que cada producto "
        "tenga cantidad suficiente. Si no hay stock, se lanza una excepción HTTP 400. La reducción de stock (RN-02) "
        "se ejecuta inmediatamente después de validar, garantizando atomicidad mediante transacciones de base de datos.")
    
    codigo_rn01 = """# Fragmento de src/modules/orders/service.py
async def create_order(order_data: OrderCreate, db: Session):
    # RN-01: Validar stock suficiente
    for detalle in order_data.detalles:
        producto = db.query(Producto).filter_by(id=detalle.producto_id).first()
        if not producto:
            raise HTTPException(404, f"Producto {detalle.producto_id} no encontrado")
        if producto.cantidad < detalle.cantidad:
            raise HTTPException(400, 
                f"Stock insuficiente para {producto.nombre}. "
                f"Disponible: {producto.cantidad}, Solicitado: {detalle.cantidad}")
    
    # RN-02: Reducir stock automáticamente
    for detalle in order_data.detalles:
        producto = db.query(Producto).filter_by(id=detalle.producto_id).first()
        producto.cantidad -= detalle.cantidad
    
    db.commit()"""
    agregar_codigo(doc, codigo_rn01)
    
    # Implementación RN-03 y RN-04
    agregar_parrafo(doc, "")
    p = doc.add_paragraph()
    p.style = 'Normal'
    run = p.add_run("Implementación de RN-03 y RN-04 (Pagos acumulativos y cambio de estado)")
    run.font.bold = True
    run.font.size = Pt(11)
    
    agregar_parrafo(doc,
        "Los pagos acumulativos (RN-03) permiten registrar múltiples pagos parciales para un mismo pedido. El campo "
        "total_pagado se actualiza sumando cada pago. El cambio automático de estado (RN-04) se implementa verificando "
        "si el monto pendiente es menor o igual a 0.01 (tolerancia para evitar problemas de precisión decimal).")
    
    codigo_rn03 = """# Fragmento de src/modules/payments/service.py
async def create_payment(payment_data: PaymentCreate, db: Session):
    pedido = db.query(Pedido).filter_by(id=payment_data.pedido_id).first()
    if not pedido:
        raise HTTPException(404, "Pedido no encontrado")
    
    # RN-03: Acumular pagos
    pedido.total_pagado += payment_data.monto
    
    # RN-04: Cambiar estado si está completamente pagado
    monto_pendiente = pedido.total - pedido.total_pagado
    if monto_pendiente <= 0.01:
        pedido.estado = "pagado"
    
    db.commit()"""
    agregar_codigo(doc, codigo_rn03)
    
    # =============================
    # 7.3 COMPONENTES PRINCIPALES
    # =============================
    agregar_titulo(doc, "7.3. Componentes principales del sistema", nivel=2)
    
    agregar_parrafo(doc,
        "El sistema se estructura en módulos funcionales siguiendo una arquitectura en capas: capa de presentación "
        "(API REST con FastAPI), capa de lógica de negocio (servicios), capa de acceso a datos (repositorios con "
        "SQLAlchemy) y capa de datos (PostgreSQL 16). A continuación se describen los componentes principales.")
    
    # Tabla 7.3: Módulos del sistema
    agregar_parrafo(doc, "")
    p_tabla3 = doc.add_paragraph()
    p_tabla3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabla3 = p_tabla3.add_run()
    run_tabla3.font.bold = True
    run_tabla3.font.size = Pt(11)
    run_tabla3.text = "Tabla 7.3. Módulos funcionales del sistema"
    
    tabla_modulos = [
        ["auth", "/api/auth/login", "Autenticación con JWT y bcrypt"],
        ["users", "/api/users/", "Gestión de usuarios internos (admin, supervisor, vendedor)"],
        ["clients", "/api/clients/", "Gestión de clientes con contactos múltiples"],
        ["products", "/api/products/", "CRUD de productos con catálogo público sin autenticación"],
        ["orders", "/api/orders/", "Creación de pedidos con validación de stock y cálculo de totales"],
        ["payments", "/api/payments/", "Registro de pagos acumulativos con actualización de estado"],
        ["devoluciones", "/api/devoluciones/", "Gestión de devoluciones con restauración de inventario"],
        ["inventory", "/api/inventory/", "Reportes de stock, alertas de stock bajo y estadísticas"]
    ]
    agregar_tabla(doc, 
                  ["Módulo", "Endpoint Base", "Funcionalidad Principal"],
                  tabla_modulos)
    
    agregar_parrafo(doc, "Fuente: Elaboración propia")
    agregar_parrafo(doc, "")
    
    agregar_parrafo(doc,
        "Cada módulo implementa el patrón MVC adaptado para APIs REST: model.py (modelos SQLAlchemy), "
        "schema.py (esquemas Pydantic para validación), service.py (lógica de negocio) y routes.py (endpoints FastAPI). "
        "Los middlewares de autenticación (JWTBearer) y auditoría (AuditMiddleware) se aplican transversalmente "
        "a todos los endpoints protegidos.")
    
    # =============================
    # 7.4 SEGURIDAD Y AUDITORÍA
    # =============================
    agregar_titulo(doc, "7.4. Implementación de seguridad y auditoría", nivel=2)
    
    agregar_parrafo(doc,
        "El sistema implementa múltiples capas de seguridad: autenticación mediante tokens JWT con expiración de "
        "30 minutos, hashing de contraseñas con bcrypt y salt aleatorio, control de acceso basado en roles (RBAC) "
        "en cada endpoint, validación de entrada con Pydantic, protección contra inyección SQL mediante ORM, y "
        "variables de entorno para credenciales sensibles (archivo .env en .gitignore). La auditoría se realiza "
        "mediante el middleware AuditMiddleware que registra automáticamente todas las acciones en logs_acciones.")
    
    # Tabla 7.4: Medidas de seguridad
    agregar_parrafo(doc, "")
    p_tabla4 = doc.add_paragraph()
    p_tabla4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tabla4 = p_tabla4.add_run()
    run_tabla4.font.bold = True
    run_tabla4.font.size = Pt(11)
    run_tabla4.text = "Tabla 7.4. Medidas de seguridad implementadas"
    
    tabla_seguridad = [
        ["Autenticación JWT", "Tokens con expiración de 30 minutos, algoritmo HS256"],
        ["Hashing de contraseñas", "bcrypt con salt aleatorio y factor de costo 12"],
        ["Control de acceso (RBAC)", "Decoradores @require_role en endpoints protegidos"],
        ["Validación de entrada", "Esquemas Pydantic con validación estricta de tipos"],
        ["Variables de entorno", "Credenciales en .env (DB_PASSWORD, SECRET_KEY) fuera del código"],
        ["Auditoría completa", "Middleware que registra usuario, endpoint, payload, IP, respuesta"]
    ]
    agregar_tabla(doc, 
                  ["Medida de Seguridad", "Implementación"],
                  tabla_seguridad)
    
    agregar_parrafo(doc, "Fuente: Elaboración propia")
    agregar_parrafo(doc, "")
    
    # CONCLUSIONES PARCIALES
    agregar_titulo(doc, "Conclusiones parciales del Epígrafe VII", nivel=2)
    
    agregar_parrafo(doc,
        "1. Se diseñó un modelo entidad-relación normalizado con 9 entidades que representan todos los conceptos "
        "del dominio, garantizando integridad referencial mediante claves primarias, foráneas y restricciones ON DELETE CASCADE.")
    
    agregar_parrafo(doc,
        "2. El esquema de base de datos implementado en PostgreSQL 16 utiliza tipos de datos optimizados (SERIAL, "
        "NUMERIC, JSONB, TIMESTAMP) y funciones almacenadas (calcular_monto_pendiente) para cálculos precisos.")
    
    agregar_parrafo(doc,
        "3. Las 9 reglas de negocio identificadas se implementaron en la capa de servicios con validaciones automáticas, "
        "manejo de transacciones y lanzamiento de excepciones HTTP descriptivas para errores.")
    
    agregar_parrafo(doc,
        "4. El sistema se estructuró en 8 módulos funcionales (auth, users, clients, products, orders, payments, "
        "devoluciones, inventory) siguiendo arquitectura en capas y patrones de diseño Repository, DTO y Dependency Injection.")
    
    agregar_parrafo(doc,
        "5. Se implementaron 6 medidas de seguridad (JWT, bcrypt, RBAC, validación Pydantic, variables de entorno, "
        "auditoría completa) que garantizan la protección de datos y trazabilidad de operaciones según RNF-05 a RNF-10.")


# Ejecutar
if __name__ == "__main__":
    print("🔄 Agregando Epígrafe VII al documento...")
    
    if not os.path.exists(FILENAME):
        print(f"❌ Error: No se encontró el archivo {FILENAME}")
        print("   Asegúrate de que el documento existe en el directorio actual.")
        exit(1)
    
    # Abrir documento existente
    doc = Document(FILENAME)
    print(f"✅ Documento cargado: {FILENAME}")
    
    # Agregar Epígrafe VII
    epigrafe_vii(doc)
    print("✅ Epígrafe VII generado")
    
    # Guardar
    doc.save(FILENAME)
    print(f"💾 Documento guardado: {FILENAME}")
    print("📝 Tus correcciones manuales anteriores se han preservado")
    print("")
    print("📊 Epígrafe VII agregado con éxito:")
    print("   - 7.1 Diseño del modelo de datos (con diagrama ER)")
    print("   - 7.1.1 Descripción de entidades y tablas")
    print("   - 7.1.2 Esquema SQL completo")
    print("   - 7.2 Implementación de reglas de negocio")
    print("   - 7.3 Componentes principales")
    print("   - 7.4 Seguridad y auditoría")
    print("   - Conclusiones parciales")
