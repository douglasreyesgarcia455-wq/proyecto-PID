"""
Script para agregar ÚNICAMENTE el Epígrafe V al documento existente
SIN regenerar todo el documento desde cero
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def epigrafe_v(doc):
    """Genera el Epígrafe V: Arquitectura y diseño de la solución propuesta"""
    # EPÍGRAFE V
    p = doc.add_paragraph('V. Arquitectura y diseño de la solución propuesta')
    # Aplicar estilo si existe
    try:
        p.style = 'Subtitulo'
    except:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(14)
    
    intro_arquitectura = """El presente epígrafe describe la arquitectura del sistema de gestión de pedidos, detallando los patrones arquitectónicos adoptados, la estructura en capas, los componentes principales y sus responsabilidades. Se presenta el diseño que garantiza la separación de responsabilidades, la mantenibilidad del código y la escalabilidad de la solución."""
    
    doc.add_paragraph(intro_arquitectura)
    
    # 5.1 Arquitectura del sistema
    doc.add_paragraph()
    p = doc.add_paragraph('5.1 Arquitectura del sistema')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_1 = """El sistema adopta una arquitectura en capas (Layered Architecture) que separa las responsabilidades en cuatro niveles fundamentales: capa de presentación, capa de servicios, capa de repositorio y capa de datos. Esta organización permite que cada capa tenga responsabilidades claramente definidas, facilitando el mantenimiento, las pruebas y la evolución del sistema."""
    
    doc.add_paragraph(texto_5_1)
    
    # 5.1.1 Capa de presentación
    doc.add_paragraph()
    p = doc.add_paragraph('5.1.1 Capa de presentación (API REST)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_1_1 = """La capa de presentación implementa una API REST mediante FastAPI, exponiendo endpoints HTTP que permiten la comunicación entre el cliente y el servidor. Esta capa se encarga de recibir las solicitudes HTTP, validar los datos de entrada mediante esquemas Pydantic, invocar los servicios de negocio correspondientes y devolver respuestas JSON estructuradas. Los controladores (routes) están organizados por módulos funcionales: autenticación, usuarios, clientes, productos, pedidos, pagos, devoluciones y reportes. La documentación automática se genera mediante OpenAPI/Swagger UI, facilitando la integración y las pruebas."""
    
    doc.add_paragraph(texto_5_1_1)
    
    # 5.1.2 Capa de servicios
    doc.add_paragraph()
    p = doc.add_paragraph('5.1.2 Capa de servicios (Lógica de negocio)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_1_2 = """La capa de servicios contiene la lógica de negocio del sistema, implementando las reglas de negocio identificadas (RN-01 a RN-09) y orquestando las operaciones entre la capa de presentación y la capa de repositorio. Los servicios son responsables de validar las reglas de negocio, gestionar transacciones, calcular totales, verificar stock, actualizar estados de pedidos, registrar logs de auditoría y coordinar operaciones complejas que involucran múltiples entidades. Esta capa garantiza que la lógica de negocio esté centralizada y sea reutilizable."""
    
    doc.add_paragraph(texto_5_1_2)
    
    # 5.1.3 Capa de repositorio
    doc.add_paragraph()
    p = doc.add_paragraph('5.1.3 Capa de repositorio (Acceso a datos)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_1_3 = """La capa de repositorio abstrae el acceso a la base de datos mediante SQLAlchemy ORM, proporcionando operaciones CRUD (Create, Read, Update, Delete) para cada entidad del modelo de datos. Los repositorios encapsulan las consultas SQL y las operaciones de persistencia, permitiendo que la capa de servicios opere con objetos Python sin conocer los detalles de implementación de la base de datos. Esta abstracción facilita las pruebas unitarias mediante mocks y permite cambiar el gestor de base de datos sin afectar las capas superiores."""
    
    doc.add_paragraph(texto_5_1_3)
    
    # 5.1.4 Capa de datos
    doc.add_paragraph()
    p = doc.add_paragraph('5.1.4 Capa de datos (PostgreSQL)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_1_4 = """La capa de datos está implementada en PostgreSQL 16, almacenando la información del sistema en nueve tablas relacionales: usuarios, clientes, contactos_clientes, productos, pedidos, detalles_pedido, pagos, devoluciones y logs_acciones. La base de datos implementa restricciones de integridad referencial mediante claves foráneas, funciones definidas por el usuario (calcular_monto_pendiente) y triggers para automatizar reglas de negocio. Las transacciones ACID garantizan la consistencia de los datos en operaciones críticas como la creación de pedidos y el registro de pagos."""
    
    doc.add_paragraph(texto_5_1_4)
    
    # 5.2 Patrones de diseño aplicados
    doc.add_paragraph()
    p = doc.add_paragraph('5.2 Patrones de diseño aplicados')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_2 = """El sistema aplica diversos patrones de diseño que mejoran la calidad del código, la mantenibilidad y la escalabilidad:"""
    
    doc.add_paragraph(texto_5_2)
    
    patrones = [
        ('Repository Pattern', 'Abstrae el acceso a datos, permitiendo que la lógica de negocio opere sin conocer los detalles de la persistencia.'),
        ('Dependency Injection', 'FastAPI inyecta automáticamente dependencias (sesiones de base de datos, servicios, usuario autenticado), facilitando las pruebas y reduciendo el acoplamiento.'),
        ('DTO (Data Transfer Object)', 'Los esquemas Pydantic actúan como DTOs, validando datos de entrada/salida y separando la representación externa de los modelos internos.'),
        ('Middleware Pattern', 'Se implementan middlewares para CORS, manejo de errores, logging de peticiones y control de acceso basado en roles (RBAC).'),
        ('Factory Pattern', 'La configuración de la base de datos y la creación de sesiones utilizan el patrón Factory para gestionar conexiones.'),
        ('Strategy Pattern', 'El sistema de autenticación puede implementar diferentes estrategias (JWT, OAuth2) mediante interfaces comunes.')
    ]
    
    for patron, descripcion in patrones:
        p = doc.add_paragraph()
        run_patron = p.add_run(f'{patron}: ')
        run_patron.bold = True
        p.add_run(descripcion)
    
    # 5.3 Seguridad en la arquitectura
    doc.add_paragraph()
    p = doc.add_paragraph('5.3 Seguridad en la arquitectura')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_3 = """La arquitectura implementa múltiples capas de seguridad para proteger la información y garantizar la integridad del sistema:"""
    
    doc.add_paragraph(texto_5_3)
    
    seguridad = [
        ('Autenticación JWT', 'Los usuarios se autentican mediante tokens JWT con expiración configurable, evitando el almacenamiento de sesiones en el servidor.'),
        ('Hashing de contraseñas', 'Las contraseñas se almacenan cifradas mediante bcrypt con salt aleatorio, garantizando que no puedan recuperarse en texto plano.'),
        ('Control de acceso basado en roles (RBAC)', 'Cada endpoint valida el rol del usuario mediante decoradores, garantizando que solo usuarios autorizados accedan a funcionalidades restringidas.'),
        ('Validación de entrada', 'Pydantic valida todos los datos de entrada, previniendo inyecciones SQL y ataques de tipo NoSQL injection.'),
        ('Variables de entorno', 'Las credenciales sensibles (secreto JWT, conexión a base de datos) se almacenan en archivos .env no versionados en Git.'),
        ('Auditoría', 'Todas las operaciones críticas se registran en la tabla logs_acciones con usuario, acción, fecha y detalles, permitiendo trazabilidad completa.')
    ]
    
    for medida, descripcion in seguridad:
        p = doc.add_paragraph()
        run_medida = p.add_run(f'{medida}: ')
        run_medida.bold = True
        p.add_run(descripcion)
    
    # 5.4 Escalabilidad y rendimiento
    doc.add_paragraph()
    p = doc.add_paragraph('5.4 Escalabilidad y rendimiento')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_5_4 = """El diseño arquitectónico considera aspectos de escalabilidad y rendimiento mediante las siguientes estrategias: uso de operaciones asíncronas en FastAPI para manejar múltiples peticiones concurrentes sin bloqueo, conexiones pooling a la base de datos mediante SQLAlchemy para reutilizar conexiones y reducir overhead, índices en columnas de búsqueda frecuente (usuarios.username, clientes.ruc, productos.nombre) para optimizar consultas, paginación en endpoints de listado para reducir la carga de datos transferidos, y separación de responsabilidades que permite escalar horizontalmente cada capa de manera independiente."""
    
    doc.add_paragraph(texto_5_4)
    
    # Conclusiones parciales del Epígrafe V
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi5 = """La adopción de una arquitectura en capas (presentación, servicios, repositorio, datos) garantiza la separación de responsabilidades, facilitando el mantenimiento, las pruebas y la evolución del sistema. La aplicación de patrones de diseño reconocidos (Repository, Dependency Injection, DTO, Middleware, Factory, Strategy) mejora la calidad del código y reduce el acoplamiento entre componentes. La implementación de múltiples capas de seguridad (JWT, bcrypt, RBAC, validación de entrada, variables de entorno, auditoría) protege la información sensible y garantiza que solo usuarios autorizados accedan a funcionalidades restringidas. Las estrategias de escalabilidad y rendimiento (operaciones asíncronas, connection pooling, índices, paginación, separación de responsabilidades) permiten que el sistema maneje cargas crecientes sin degradación del desempeño. En conjunto, la arquitectura propuesta establece una base sólida para un sistema robusto, seguro, mantenible y escalable, alineado con las mejores prácticas de desarrollo de software empresarial."""
    
    doc.add_paragraph(conclusiones_epi5)
    
    # Salto de página
    doc.add_page_break()

def main():
    filename = 'Informe_Tecnico_PID_NUEVO_V1.docx'
    
    print(f"📄 Abriendo documento existente: {filename}")
    doc = Document(filename)
    
    print("➕ Agregando Epígrafe V al final del documento...")
    epigrafe_v(doc)
    
    print("💾 Guardando cambios...")
    doc.save(filename)
    
    print(f"\n✅ Epígrafe V agregado exitosamente a: {filename}")
    print("📝 Tus correcciones manuales anteriores se han preservado.")

if __name__ == "__main__":
    main()
