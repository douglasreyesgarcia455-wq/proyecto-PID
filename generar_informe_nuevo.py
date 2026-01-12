"""
Generador de Informe Técnico - NUEVO DOCUMENTO
Sistema de Gestión de Pedidos - Proyecto PID III

Este script genera un nuevo documento DOCX desde cero con las primeras secciones
según la Guía UCI, hasta el desarrollo (sin incluirlo aún).

Epígrafes a incluir en este paso:
1. Portada
2. Resumen (español)
3. Abstract (inglés)
4. Índices (General, Tablas, Figuras) - Placeholders
5. Opinión del Tutor - Placeholder
6. Introducción

Instrucciones:
- Generar epígrafe por epígrafe
- Crear todas las tablas e información paso a paso
- Referencias en APA 7ma (archivo separado: REFERENCIAS_BIBLIOGRAFICAS.md)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def crear_estilos(doc):
    """Crea estilos personalizados para el documento"""
    styles = doc.styles
    
    # Estilo para títulos de capítulo
    if 'Titulo Capitulo' not in styles:
        titulo_cap = styles.add_style('Titulo Capitulo', WD_STYLE_TYPE.PARAGRAPH)
        titulo_cap.font.name = 'Arial'
        titulo_cap.font.size = Pt(16)
        titulo_cap.font.bold = True
        titulo_cap.font.color.rgb = RGBColor(0, 0, 0)
        titulo_cap.paragraph_format.space_before = Pt(24)
        titulo_cap.paragraph_format.space_after = Pt(12)
        titulo_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Estilo para subtítulos
    if 'Subtitulo' not in styles:
        subtitulo = styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        subtitulo.font.name = 'Arial'
        subtitulo.font.size = Pt(14)
        subtitulo.font.bold = True
        subtitulo.paragraph_format.space_before = Pt(12)
        subtitulo.paragraph_format.space_after = Pt(6)
    
    # Estilo para texto normal
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.line_spacing = 1.5

def portada(doc):
    """Genera la portada del documento"""
    # Universidad
    p = doc.add_paragraph('UNIVERSIDAD DE LAS CIENCIAS INFORMÁTICAS')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    
    # Facultad
    p = doc.add_paragraph('FACULTAD DE TECNOLOGÍAS INTERACTIVAS')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Tipo de trabajo
    p = doc.add_paragraph('Trabajo de Curso')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    
    # Asignatura
    p = doc.add_paragraph('Proyecto Integrador de Desarrollo III')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Título (destacado)
    p = doc.add_paragraph('Sistema de Gestión de Pedidos para MIPYME')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('Comercializadora de Ácido Acético y Botellas Plásticas')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Autores
    p = doc.add_paragraph('Autores:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('Douglas Reyes García')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    p = doc.add_paragraph('Alex Daniel Jorro Gacita')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    # Espacio
    doc.add_paragraph()
    
    # Tutor
    p = doc.add_paragraph('Tutor:')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    
    p = doc.add_paragraph('Lisset Salazar Gómez')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    # Espacio
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Fecha y lugar
    p = doc.add_paragraph('La Habana, Enero 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    
    p = doc.add_paragraph('"Año 67 de la Revolución"')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(10)
    p.runs[0].font.italic = True
    
    # Salto de página
    doc.add_page_break()

def resumen(doc):
    """Genera el resumen en español"""
    # Título
    p = doc.add_paragraph('Resumen')
    p.style = 'Titulo Capitulo'
    
    # Contenido del resumen (del documento ISW_MYPIME)
    texto_resumen = """El presente trabajo aborda la propuesta de un Sistema de Gestión de Pedidos para una MIPYME comercializadora de ácido acético y botellas plásticas, orientado a optimizar el control de sus operaciones comerciales mediante la automatización de los procesos de pedidos, pagos, entregas y devoluciones. Actualmente, la empresa enfrenta dificultades derivadas del manejo manual de la información, lo que genera demoras, errores y falta de trazabilidad en la gestión administrativa y operativa.

El objetivo general del trabajo consiste en analizar la factibilidad y diseñar la base conceptual de una solución que centralice la información del negocio, permitiendo gestionar clientes, productos y transacciones de forma eficiente, confiable y accesible. La propuesta busca garantizar la integridad de los datos, reducir la carga operativa del personal y facilitar la toma de decisiones mediante reportes automatizados.

Para el desarrollo conceptual del sistema se adopta un enfoque ágil, basado en la metodología Extreme Programming (XP), por su adecuación a equipos pequeños, requisitos parcialmente cambiantes y necesidad de entregas funcionales continuas. Este enfoque favorece la retroalimentación constante del cliente y la mejora progresiva del producto.

El trabajo se centra en el estudio lógico del problema, la identificación de los procesos de negocio y la fundamentación teórica necesaria para etapas posteriores de diseño e implementación del sistema, contribuyendo así al fortalecimiento de la gestión comercial en pequeñas empresas del sector químico."""
    
    p = doc.add_paragraph(texto_resumen)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Palabras clave
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Palabras clave: ').bold = True
    p.add_run('MIPYME; Gestión de Pedidos; Ingeniería de Software; Automatización Comercial; XP; FastAPI; PostgreSQL; Python.')
    
    # Salto de página
    doc.add_page_break()

def abstract(doc):
    """Genera el abstract en inglés"""
    # Título
    p = doc.add_paragraph('Abstract')
    p.style = 'Titulo Capitulo'
    
    # Contenido del abstract (traducción)
    texto_abstract = """This work addresses the proposal for an Order Management System for a MIPYME (micro, small and medium-sized enterprise) that commercializes acetic acid and plastic bottles, aimed at optimizing the control of its commercial operations through the automation of order, payment, delivery, and return processes. Currently, the company faces difficulties derived from manual information handling, which generates delays, errors, and lack of traceability in administrative and operational management.

The general objective of the work is to analyze the feasibility and design the conceptual basis of a solution that centralizes business information, allowing efficient, reliable, and accessible management of clients, products, and transactions. The proposal seeks to guarantee data integrity, reduce operational burden on staff, and facilitate decision-making through automated reports.

For the conceptual development of the system, an agile approach is adopted, based on the Extreme Programming (XP) methodology, due to its suitability for small teams, partially changing requirements, and the need for continuous functional deliveries. This approach favors constant customer feedback and progressive product improvement.

The work focuses on the logical study of the problem, the identification of business processes, and the theoretical foundation necessary for subsequent stages of system design and implementation, thus contributing to the strengthening of commercial management in small companies in the chemical sector."""
    
    p = doc.add_paragraph(texto_abstract)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Keywords
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Keywords: ').bold = True
    p.add_run('MIPYME; Order Management; Software Engineering; Commercial Automation; XP; FastAPI; PostgreSQL; Python.')
    
    # Salto de página
    doc.add_page_break()

def indices_placeholders(doc):
    """Genera placeholders para índices (deben generarse automáticamente en Word)"""
    # Índice General
    p = doc.add_paragraph('Índice General')
    p.style = 'Titulo Capitulo'
    
    p = doc.add_paragraph('[El índice general se generará automáticamente en Microsoft Word]')
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph('Instrucciones:')
    p.runs[0].bold = True
    
    instrucciones = [
        '1. Abrir el documento en Microsoft Word',
        '2. Hacer clic donde debe ir el índice',
        '3. Ir a la pestaña "Referencias"',
        '4. Hacer clic en "Tabla de contenido"',
        '5. Seleccionar "Tabla automática 1" o "Tabla automática 2"'
    ]
    
    for instr in instrucciones:
        doc.add_paragraph(instr, style='List Number')
    
    doc.add_page_break()
    
    # Índice de Tablas
    p = doc.add_paragraph('Índice de Tablas')
    p.style = 'Titulo Capitulo'
    
    p = doc.add_paragraph('[El índice de tablas se generará automáticamente en Microsoft Word]')
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph('Instrucciones:')
    p.runs[0].bold = True
    
    instrucciones_tablas = [
        '1. Abrir el documento en Microsoft Word',
        '2. Hacer clic donde debe ir el índice de tablas',
        '3. Ir a la pestaña "Referencias"',
        '4. Hacer clic en "Insertar tabla de ilustraciones"',
        '5. En "Etiqueta de título" seleccionar "Tabla"',
        '6. Hacer clic en "Aceptar"'
    ]
    
    for instr in instrucciones_tablas:
        doc.add_paragraph(instr, style='List Number')
    
    doc.add_page_break()
    
    # Índice de Figuras
    p = doc.add_paragraph('Índice de Figuras')
    p.style = 'Titulo Capitulo'
    
    p = doc.add_paragraph('[El índice de figuras se generará automáticamente en Microsoft Word]')
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    p = doc.add_paragraph('Instrucciones:')
    p.runs[0].bold = True
    
    instrucciones_figuras = [
        '1. Abrir el documento en Microsoft Word',
        '2. Hacer clic donde debe ir el índice de figuras',
        '3. Ir a la pestaña "Referencias"',
        '4. Hacer clic en "Insertar tabla de ilustraciones"',
        '5. En "Etiqueta de título" seleccionar "Figura"',
        '6. Hacer clic en "Aceptar"'
    ]
    
    for instr in instrucciones_figuras:
        doc.add_paragraph(instr, style='List Number')
    
    doc.add_page_break()

def opinion_tutor_placeholder(doc):
    """Genera placeholder para opinión del tutor"""
    p = doc.add_paragraph('Opinión del Tutor')
    p.style = 'Titulo Capitulo'
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('[PÁGINA RESERVADA PARA LA OPINIÓN DEL TUTOR]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph('Esta página debe ser completada por el tutor del proyecto.')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].italic = True
    p.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()

def introduccion(doc):
    """Genera el capítulo de Introducción (del documento ISW_MYPIME)"""
    # Título
    p = doc.add_paragraph('Introducción')
    p.style = 'Titulo Capitulo'
    
    # Contexto y problemática
    parrafo1 = """La gestión de pedidos constituye una de las actividades más relevantes en las micro, pequeñas y medianas empresas (MIPYMES) dedicadas a la comercialización de productos químicos, donde la eficiencia operativa, la exactitud en el registro de las transacciones y la trazabilidad del proceso resultan esenciales para garantizar la satisfacción del cliente y la rentabilidad del negocio. En la comercializadora de ácido acético objeto de estudio, este proceso se realiza actualmente mediante procedimientos manuales que involucran el uso de hojas de cálculo y comunicación informal entre clientes, vendedores y personal administrativo."""
    
    doc.add_paragraph(parrafo1)
    
    parrafo2 = """Dichas prácticas generan demoras en el procesamiento de los pedidos, errores en la consolidación de la información, dificultades en el control del inventario y poca transparencia en el seguimiento de las ventas. Este escenario limita la capacidad de respuesta de la empresa, afecta la confiabilidad de los registros y obstaculiza la toma de decisiones basada en datos actualizados. La falta de trazabilidad y de una gestión centralizada también dificulta el control de las devoluciones y pagos, elementos fundamentales para mantener la integridad operativa del negocio."""
    
    doc.add_paragraph(parrafo2)
    
    # Planteamiento del problema
    doc.add_paragraph()
    p = doc.add_paragraph('Planteamiento del Problema')
    p.style = 'Subtitulo'
    
    parrafo3 = """A partir de esta situación, se plantea el siguiente problema de investigación: ¿Cómo automatizar la gestión de pedidos de una MIPYME comercializadora de ácido acético y botellas plásticas, mejorando la eficiencia, trazabilidad y calidad del servicio?"""
    
    doc.add_paragraph(parrafo3)
    
    # Objetivo General
    doc.add_paragraph()
    p = doc.add_paragraph('Objetivo General')
    p.style = 'Subtitulo'
    
    parrafo4 = """Diseñar e implementar un sistema web de gestión de pedidos para la comercializadora de ácido acético y botellas plásticas, que permita automatizar los procesos de solicitud, validación, registro y seguimiento de pedidos, garantizando la coherencia de la información y la mejora del desempeño organizacional."""
    
    doc.add_paragraph(parrafo4)
    
    # Tareas de Investigación
    doc.add_paragraph()
    p = doc.add_paragraph('Tareas de Investigación')
    p.style = 'Subtitulo'
    
    tareas = [
        'Analizar el proceso actual de gestión de pedidos en la empresa.',
        'Identificar los requerimientos funcionales y no funcionales del sistema.',
        'Modelar los procesos de negocio mediante técnicas de ingeniería de software.',
        'Diseñar la arquitectura del software y la estructura de la base de datos.',
        'Implementar los módulos de la aplicación mediante un enfoque iterativo e incremental basado en Extreme Programming (XP).',
        'Validar el sistema a través de pruebas técnicas y de usabilidad.'
    ]
    
    for tarea in tareas:
        doc.add_paragraph(tarea, style='List Bullet')
    
    # Metodología
    doc.add_paragraph()
    p = doc.add_paragraph('Metodología de Investigación')
    p.style = 'Subtitulo'
    
    parrafo5 = """Metodológicamente, la investigación combina métodos teóricos, empíricos y técnicos. Entre los métodos teóricos se emplearon el analítico-sintético y el inductivo-deductivo, utilizados para el estudio del estado del arte y la definición de los requerimientos. Los métodos empíricos incluyeron la observación directa del proceso comercial, entrevistas con los empleados y revisión documental de los registros de pedidos existentes."""
    
    doc.add_paragraph(parrafo5)
    
    parrafo6 = """Además, se aplicaron herramientas propias de la Ingeniería de Software, como el modelado UML para la representación estructural y comportamental del sistema, el uso de herramientas CASE (Lucidchart y Visual Paradigm) para la elaboración de diagramas, y la utilización de tecnologías web modernas: Python 3.12 con FastAPI para el backend, HTML5, CSS3 y JavaScript ES2023 para la interfaz de usuario, y PostgreSQL 16 como gestor de base de datos. Estas tecnologías garantizan la integridad, escalabilidad y mantenibilidad del sistema desarrollado."""
    
    doc.add_paragraph(parrafo6)
    
    # Estructura del Documento
    doc.add_paragraph()
    p = doc.add_paragraph('Estructura del Documento')
    p.style = 'Subtitulo'
    
    parrafo7 = """El presente informe se estructura en ocho capítulos que abarcan desde la fundamentación teórica hasta la validación del sistema implementado. El Capítulo I establece el marco conceptual y tecnológico del proyecto. El Capítulo II modela el contexto del negocio y las reglas que lo rigen. El Capítulo III documenta los requisitos funcionales y no funcionales. El Capítulo IV describe la validación y gestión de requisitos mediante casos de prueba. Los capítulos V y VI presentan el modelado de la estructura, comportamiento y diseño del sistema. El Capítulo VII detalla la implementación técnica, y el Capítulo VIII expone los resultados de las pruebas de validación. Finalmente, se presentan las conclusiones, recomendaciones y referencias bibliográficas."""
    
    doc.add_paragraph(parrafo7)
    
    # Salto de página
    doc.add_page_break()

def desarrollo(doc):
    """Genera el capítulo DESARROLLO con introducción integrada y Epígrafe I"""
    # Título
    p = doc.add_paragraph('DESARROLLO')
    p.style = 'Titulo Capitulo'
    
    # Introducción integrada (sin título de sección)
    intro = """El presente capítulo desarrolla el proceso de diseño e implementación del sistema web de gestión de pedidos para la MIPYME comercializadora de ácido acético y botellas plásticas. Se estructura en ocho epígrafes que abarcan desde la fundamentación teórica hasta la validación del sistema implementado. Cada epígrafe contribuye a demostrar la pertinencia de la solución propuesta, la adecuada selección tecnológica y metodológica, y el cumplimiento de los objetivos planteados."""
    
    doc.add_paragraph(intro)
    doc.add_paragraph()
    
    # EPÍGRAFE I
    p = doc.add_paragraph('I. Gestión de pedidos y automatización de procesos comerciales en MIPYMES')
    p.style = 'Subtitulo'
    
    # 1.1 Conceptos asociados al tema
    doc.add_paragraph()
    p = doc.add_paragraph('1.1 Conceptos asociados al tema')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # 1.1.1
    doc.add_paragraph()
    p = doc.add_paragraph('1.1.1 Gestión comercial en las MIPYMES')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_1_1_1 = """La gestión comercial es el conjunto de procesos administrativos, operativos y estratégicos orientados a planificar, ejecutar y controlar actividades de venta, distribución y atención al cliente, utilizando herramientas que centralicen información de clientes, productos, pedidos y pagos para mejorar la toma de decisiones, la eficiencia operativa y la competitividad (Ramírez, 2021)."""
    
    doc.add_paragraph(texto_1_1_1)
    
    # 1.1.2
    doc.add_paragraph()
    p = doc.add_paragraph('1.1.2 Automatización de procesos empresariales')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_1_1_2 = """La automatización de procesos empresariales consiste en la capacidad de los sistemas informáticos para ejecutar tareas rutinarias sin intervención manual directa, garantizando rapidez, precisión y trazabilidad; en las MIPYMES permite reducir errores, optimizar tiempos de gestión, mantener actualizados los registros de pedidos, pagos e inventarios y destinar más tiempo a actividades estratégicas (González & Pérez, 2022)."""
    
    doc.add_paragraph(texto_1_1_2)
    
    # 1.1.3
    doc.add_paragraph()
    p = doc.add_paragraph('1.1.3 Gestión de relaciones con el cliente (CRM)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_1_1_3 = """La gestión de relaciones con el cliente (CRM) es un modelo apoyado en tecnologías de la información que facilita organizar, analizar y mejorar las interacciones con los clientes mediante la integración de contactos, historial de compras, comunicaciones y transacciones, permitiendo estrategias de fidelización, personalización del servicio y una visión integral del desempeño comercial en las MIPYMES (López, 2020)."""
    
    doc.add_paragraph(texto_1_1_3)
    
    # 1.1.4
    doc.add_paragraph()
    p = doc.add_paragraph('1.1.4 Integración de gestión comercial, automatización y CRM')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_1_1_4 = """La integración de estos elementos constituye una base esencial para el desarrollo competitivo y sostenible de las MIPYMES, ya que contribuye a modernizar las operaciones, aumentar la eficiencia y fortalecer la satisfacción del cliente (Ramírez, 2021; González & Pérez, 2022; López, 2020)."""
    
    doc.add_paragraph(texto_1_1_4)
    
    # Conclusiones parciales del Epígrafe I
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi1 = """El estudio de los fundamentos teóricos relacionados con la gestión comercial, la automatización de procesos y los sistemas CRM evidencia la importancia de centralizar la información empresarial para mejorar la eficiencia operativa y la toma de decisiones en las MIPYMES. La integración de estos conceptos constituye el marco teórico que fundamenta el desarrollo de un sistema de gestión de pedidos adaptado a las necesidades específicas de pequeñas empresas del sector químico, orientado a superar las limitaciones derivadas de procesos manuales y a fortalecer la competitividad del negocio."""
    
    doc.add_paragraph(conclusiones_epi1)
    
    # Salto de página
    doc.add_page_break()

def epigrafe_ii(doc):
    """Genera el Epígrafe II: Sistemas de información para la gestión de pedidos"""
    # EPÍGRAFE II
    p = doc.add_paragraph('II. Sistemas de información para la gestión de pedidos')
    p.style = 'Subtitulo'
    
    # 2.1 Análisis de mercado
    doc.add_paragraph()
    p = doc.add_paragraph('2.1 Análisis de mercado')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    intro_analisis = """El estudio de soluciones existentes revela diversas plataformas CRM orientadas a la gestión de clientes y procesos comerciales en pequeñas y medianas empresas. La Tabla 2.1 presenta el análisis comparativo de sistemas relevantes que sirven como referencia para el desarrollo de la solución propuesta."""
    
    doc.add_paragraph(intro_analisis)
    doc.add_paragraph()
    
    # Tabla comparativa
    p = doc.add_paragraph('Tabla 2.1: Análisis comparativo de plataformas de gestión comercial')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Crear tabla con 7 filas (1 encabezado + 6 sistemas) y 5 columnas
    table = doc.add_table(rows=7, cols=5)
    table.style = 'Light Grid Accent 1'
    
    # Configurar encabezados
    headers = ['Sistema', 'Tipo', 'Gestión de Clientes', 'Gestión de Pedidos y Pagos', 'Enfoque']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Aplicar formato bold a los encabezados
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Datos de los sistemas
    sistemas = [
        ['HubSpot CRM', 'CRM en la nube', 'Sí', 'Parcial', 'Marketing y ventas digitales'],
        ['Zoho CRM', 'SaaS Empresarial', 'Sí', 'Sí', 'Ventas y automatización comercial'],
        ['Bitrix24', 'CRM colaborativo', 'Sí', 'Sí', 'Comunicación interna y ventas'],
        ['Odoo CRM', 'Módulo integrado', 'Sí', 'Sí', 'Gestión integral empresarial'],
        ['Freshsales', 'CRM en la nube', 'Sí', 'Parcial', 'Seguimiento de clientes potenciales'],
        ['Zoom LC Odoo', 'ERP y CRM', 'Sí', 'Sí', 'Gestión integral empresarial con módulos de ventas, compras, contabilidad y atención al cliente (adaptado por DESOFT al entorno cubano)']
    ]
    
    # Rellenar datos
    for i, sistema in enumerate(sistemas, start=1):
        row_cells = table.rows[i].cells
        for j, dato in enumerate(sistema):
            row_cells[j].text = dato
            # Aplicar formato
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Identificación de brechas
    gap_analysis = """El análisis identifica como puntos coincidentes la centralización de la información de clientes y la automatización de procesos de venta. Sin embargo, se evidencia una brecha en soluciones personalizadas para MIPYMES locales que combinen la gestión simplificada de pedidos, pagos, productos y devoluciones, en un entorno accesible, ligero y adaptado a los recursos tecnológicos limitados de este tipo de negocios."""
    
    doc.add_paragraph(gap_analysis)
    
    # Justificación de la solución personalizada
    doc.add_paragraph()
    justificacion = """La necesidad de una solución personalizada se justifica por la especificidad del contexto operativo de las MIPYMES cubanas: recursos limitados, conectividad variable, procesos comerciales particulares y requisitos de trazabilidad que no son adecuadamente cubiertos por las plataformas genéricas de CRM disponibles en el mercado internacional. Además, la dependencia de servicios en la nube (SaaS) implica costos recurrentes de suscripción y requerimientos de conectividad estable que no siempre están disponibles en el entorno local."""
    
    doc.add_paragraph(justificacion)
    
    # Conclusiones parciales del Epígrafe II
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi2 = """El análisis comparativo de las soluciones CRM existentes en el mercado evidencia que, aunque estas plataformas ofrecen funcionalidades robustas para la gestión de clientes y procesos comerciales, no responden de manera adecuada a las necesidades específicas de las MIPYMES cubanas. La identificación de esta brecha justifica el desarrollo de una solución personalizada, adaptada al contexto local, que integre de forma ligera y eficiente la gestión de pedidos, pagos, productos, clientes y devoluciones, minimizando la dependencia de conectividad externa y ajustándose a los recursos tecnológicos disponibles en este tipo de organizaciones."""
    
    doc.add_paragraph(conclusiones_epi2)
    
    # Salto de página
    doc.add_page_break()

def epigrafe_iii(doc):
    """Genera el Epígrafe III: Diagnóstico del proceso de gestión de pedidos en la MIPYME"""
    # EPÍGRAFE III
    p = doc.add_paragraph('III. Diagnóstico del proceso de gestión de pedidos en la MIPYME')
    p.style = 'Subtitulo'
    
    intro_diagnostico = """El presente epígrafe tiene como objetivo analizar el modelo de negocio actual de la MIPYME comercializadora de ácido acético y botellas plásticas, identificando los procesos, actores y limitaciones del sistema existente. A través de técnicas de modelado empresarial y recopilación de información, se caracteriza el contexto operativo de la empresa, evidenciando las oportunidades de mejora que justifican el desarrollo de una solución tecnológica adaptada a sus necesidades específicas."""
    
    doc.add_paragraph(intro_diagnostico)
    
    # 3.1 Técnicas de recopilación de información
    doc.add_paragraph()
    p = doc.add_paragraph('3.1 Técnicas de recopilación de información')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_3_1 = """Para la obtención de información se emplearon diversas técnicas cualitativas orientadas a comprender la dinámica del negocio y los procesos vinculados a la gestión comercial. En primer lugar, se desarrolló una sesión de lluvia de ideas entre los integrantes del equipo, con el propósito de identificar las principales dificultades en la administración de pedidos, pagos y clientes, así como las oportunidades de mejora en la organización de la información. Posteriormente, se realizaron entrevistas semiestructuradas que permitieron profundizar en el flujo actual de trabajo, detectar ineficiencias en la comunicación interna y conocer las percepciones de los participantes sobre el desempeño de las actividades comerciales. De forma complementaria, se aplicó la observación directa durante la ejecución de las tareas operativas, lo que permitió evidenciar la dependencia de procesos manuales y la falta de centralización en el manejo de datos. Finalmente, se efectuó un análisis documental de registros y documentos internos relacionados con pedidos, facturación y pagos, con el fin de evaluar la trazabilidad, consistencia y disponibilidad de la información existente. La aplicación de estas técnicas permitió obtener una comprensión integral del contexto operativo y definir las principales necesidades y oportunidades de mejora en la gestión comercial."""
    
    doc.add_paragraph(texto_3_1)
    
    # 3.2 Fuentes de obtención de requisitos
    doc.add_paragraph()
    p = doc.add_paragraph('3.2 Fuentes de obtención de requisitos')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_3_2 = """En el proceso de obtención de requisitos, es fundamental identificar las diversas fuentes de información que permiten definir con precisión las necesidades del sistema. Entre estas fuentes se consideran las metas del negocio, los stakeholders o interesados, el conocimiento del dominio, el entorno operacional y el entorno organizacional, cada una aportando perspectivas complementarias. Las metas u objetivos del negocio permiten identificar los resultados específicos que el sistema debe apoyar o mejorar, como reducir el tiempo de atención al cliente, optimizar la gestión de inventario o incrementar la eficiencia de los procesos de pago; para obtener esta información se analizaron planes estratégicos, informes de gestión y objetivos de mejora definidos por la dirección, lo que facilitó priorizar los requisitos y asegurar que estuvieran alineados con las necesidades reales de la organización. Los stakeholders, por su parte, incluyen a todas las personas o grupos con interés, responsabilidad o influencia sobre el sistema, y sus necesidades, expectativas y restricciones constituyen una fuente directa de requisitos; en este caso, se emplearon entrevistas, encuestas y tallares con usuarios finales, clientes, patrocinadores y personal operativo, asegurando así que se capturaran tanto los requerimientos funcionales como los no funcionales."""
    
    doc.add_paragraph(texto_3_2)
    
    doc.add_paragraph()
    texto_3_2_cont = """El conocimiento del dominio aporta información sobre los procesos, reglas de negocio y prácticas del sector, permitiendo comprender cómo debe comportarse el sistema y qué restricciones deben considerarse; este conocimiento se obtuvo mediante la revisión de manuales, normativas de la industria, documentación interna y reuniones con expertos del área, garantizando que los requisitos reflejaran las realidades del negocio. Además, el entorno operacional considera las condiciones físicas, técnicas y ambientales en las que funcionará el sistema, incluyendo factores como ubicación, conectividad, dispositivos y equipos de soporte; se recopiló mediante inspecciones de campo, análisis de infraestructura tecnológica y reuniones con el personal operativo, con el fin de definir las restricciones y necesidades técnicas críticas. Finalmente, el entorno organizacional abarca la cultura, estructura, procesos internos y normas de la empresa, como jerarquías, protocolos y políticas de seguridad, que influyen directamente en la viabilidad y adaptación del sistema; para capturar esta información se realizaron entrevistas con líderes de área, revisión de manuales organizacionales y análisis de los procesos existentes, asegurando que los requisitos se ajustaran a las prácticas y políticas internas de la organización."""
    
    doc.add_paragraph(texto_3_2_cont)
    
    # 3.3 Modelo conceptual del negocio
    doc.add_paragraph()
    p = doc.add_paragraph('3.3 Modelo conceptual del negocio')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_3_3 = """El modelo conceptual representa las principales entidades y relaciones que intervienen en la gestión comercial del negocio, abarcando desde la atención al cliente hasta el control administrativo de pagos y devoluciones. En este modelo, el proceso inicia con el cliente, quien realiza pedidos de productos. Cada pedido puede estar compuesto por uno o varios productos y constituye el punto de partida del ciclo comercial. El cliente también puede realizar devoluciones de un pedido; en caso de que se devuelva solo una parte, el pedido se modifica para reflejar la cantidad restante."""
    
    doc.add_paragraph(texto_3_3)
    
    doc.add_paragraph()
    texto_3_3_cont1 = """El cliente es atendido por un vendedor, responsable de gestionar las ventas correspondientes a los pedidos. Cada venta puede implicar uno o varios pagos, dependiendo del valor total del pedido. El supervisor cumple la función de vigilar el desempeño del vendedor y supervisar el avance de las ventas, asegurando el cumplimiento de los procedimientos establecidos."""
    
    doc.add_paragraph(texto_3_3_cont1)
    
    doc.add_paragraph()
    texto_3_3_cont2 = """Cuando un vendedor concreta una venta, registra la información asociada, como la fecha y el responsable de la operación. Estos datos se almacenan en un historial controlado por el administrador, quien también tiene a su cargo la verificación y control de los pagos. Los pagos generan información complementaria que se guarda en el historial junto con los detalles de la venta y, además, cada pago genera una factura que respalda formalmente la transacción."""
    
    doc.add_paragraph(texto_3_3_cont2)
    
    doc.add_paragraph()
    texto_3_3_cont3 = """Asimismo, las devoluciones generan registros que se documentan en el historial, permitiendo mantener un control actualizado de las operaciones. Los productos son administrados exclusivamente por el administrador, quien gestiona su disponibilidad, actualiza los movimientos relacionados con ventas o devoluciones y vela por la coherencia de la información registrada en el sistema comercial."""
    
    doc.add_paragraph(texto_3_3_cont3)
    
    doc.add_paragraph()
    texto_3_3_concl = """En conjunto, el modelo refleja una estructura organizada que permite comprender de manera integral el flujo de información y la interacción entre los distintos actores del negocio, asegurando el control y seguimiento de las operaciones en todas las etapas del proceso comercial."""
    
    doc.add_paragraph(texto_3_3_concl)
    
    # 3.4 Reglas del negocio
    doc.add_paragraph()
    p = doc.add_paragraph('3.4 Reglas del negocio')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    intro_reglas = """A continuación se especifican las nueve reglas que rigen el proceso de gestión comercial de la MIPYME, las cuales fueron identificadas durante el análisis del modelo de negocio y constituyen restricciones y políticas operativas que el sistema debe implementar:"""
    
    doc.add_paragraph(intro_reglas)
    doc.add_paragraph()
    
    # Tabla de reglas de negocio
    p = doc.add_paragraph('Tabla 3.1: Reglas del negocio')
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Crear tabla con 10 filas (1 encabezado + 9 reglas) y 3 columnas
    table = doc.add_table(rows=10, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Configurar encabezados
    headers = ['No.', 'Regla', 'Descripción']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
    
    # Datos de las reglas
    reglas = [
        ['RN-01', 'Validación de stock antes de crear pedido', 'El sistema debe verificar que exista stock suficiente de cada producto antes de permitir la creación de un pedido.'],
        ['RN-02', 'Reducción automática de stock', 'Al confirmar un pedido, el sistema reduce automáticamente la cantidad en inventario de los productos incluidos.'],
        ['RN-03', 'Pagos acumulativos por pedido', 'Un pedido puede recibir múltiples pagos parciales que se acumulan hasta completar el total.'],
        ['RN-04', 'Cambio automático de estado a "pagado"', 'Cuando la suma de pagos de un pedido alcanza o supera el total, el sistema cambia automáticamente su estado a "pagado".'],
        ['RN-05', 'Cálculo exacto de monto pendiente', 'El sistema calcula el monto pendiente de cada pedido mediante una función SQL: total_pedido - suma_pagos.'],
        ['RN-06', 'Restricción de sobrepago', 'No se permite registrar un pago que exceda el monto pendiente del pedido.'],
        ['RN-07', 'Devolución cambia estado del pedido', 'Al registrar una devolución, el estado del pedido cambia automáticamente a "devuelto".'],
        ['RN-08', 'Restauración de inventario en devolución', 'Al registrar una devolución, el sistema restaura automáticamente las cantidades devueltas al inventario.'],
        ['RN-09', 'Eliminación de pagos en devolución', 'Al registrar una devolución de pedido, todos los pagos asociados se eliminan automáticamente.']
    ]
    
    # Rellenar datos
    for i, regla in enumerate(reglas, start=1):
        row_cells = table.rows[i].cells
        for j, dato in enumerate(regla):
            row_cells[j].text = dato
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    
    doc.add_paragraph()
    
    # Conclusiones parciales del Epígrafe III
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi3 = """El análisis del modelo de negocio actual permitió identificar los roles, actores y reglas que intervienen en los procesos de la MIPYME, poniendo en evidencia las limitaciones derivadas del manejo manual de pedidos, pagos, devoluciones e inventario. La aplicación de técnicas de recopilación de información (lluvia de ideas, entrevistas semiestructuradas, observación directa y análisis documental) facilitó la comprensión integral del contexto operativo y las necesidades reales de los usuarios. La identificación de las fuentes de obtención de requisitos (stakeholders, metas del negocio, conocimiento del dominio, entorno operacional y organizacional) garantizó que los requisitos del sistema estuvieran alineados con las prácticas y políticas internas de la organización. La elaboración del modelo conceptual permitió comprender las interacciones entre los actores del negocio y las actividades principales, revelando redundancias y oportunidades de automatización. Finalmente, la especificación de las nueve reglas del negocio proporcionó una base sólida para el diseño posterior del sistema, asegurando la coherencia entre los procesos reales y las mejoras planteadas, y contribuyendo al fortalecimiento de la gestión comercial de la MIPYME."""
    
    doc.add_paragraph(conclusiones_epi3)
    
    # Salto de página
    doc.add_page_break()

def epigrafe_iv(doc):
    """Genera el Epígrafe IV: Metodología, tecnologías y herramientas para el desarrollo del sistema"""
    # EPÍGRAFE IV
    p = doc.add_paragraph('IV. Metodología, tecnologías y herramientas para el desarrollo del sistema')
    p.style = 'Subtitulo'
    
    intro_metodologia = """El presente epígrafe tiene como objetivo sistematizar y justificar las decisiones tecnológicas y metodológicas adoptadas para el desarrollo del sistema de gestión de pedidos. Se establece el marco metodológico basado en Extreme Programming (XP), se fundamenta el enfoque ágil mediante el modelo de Boehm y Turner, y se describen las herramientas y tecnologías seleccionadas, garantizando la coherencia entre las características del proyecto y los recursos disponibles."""
    
    doc.add_paragraph(intro_metodologia)
    
    # 4.1 Fundamentación del proceso de software a desarrollar
    doc.add_paragraph()
    p = doc.add_paragraph('4.1 Fundamentación del proceso de software a desarrollar')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_1 = """El contexto empresarial en el que opera una MIPYME exige un manejo confiable y oportuno de la información, ya que las actividades comerciales dependen de la disponibilidad inmediata de datos precisos para la gestión de pedidos, pagos y control de inventarios. Este entorno requiere procesos organizativos capaces de asegurar integridad y coherencia en los registros, así como trazabilidad en cada operación realizada. Además, la dinámica del negocio demanda la capacidad de adaptarse rápidamente a cambios en las políticas comerciales, volúmenes de trabajo o necesidades específicas de los clientes, garantizando siempre continuidad operativa y una respuesta eficiente en cada interacción."""
    
    doc.add_paragraph(texto_4_1)
    
    # 4.1.1 Enfoque de ingeniería de software
    doc.add_paragraph()
    p = doc.add_paragraph('4.1.1 Enfoque de ingeniería de software')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_1_1 = """La aplicación del método de la estrella de Boehm y Turner permite determinar la idoneidad de un enfoque ágil basado en las características del proyecto. La evaluación posiciona este desarrollo cerca del centro del modelo, debido a los siguientes factores: un equipo reducido de desarrolladores (2 integrantes), requisitos parcialmente estables, criticidad moderada y necesidad de entregas tempranas y funcionales. Esta combinación favorece un proceso iterativo e incremental, centrado en la comunicación constante con el cliente y la mejora continua del producto."""
    
    doc.add_paragraph(texto_4_1_1)
    
    doc.add_paragraph()
    texto_caracteristicas = """El proyecto presenta las características siguientes:"""
    p = doc.add_paragraph(texto_caracteristicas)
    p.runs[0].bold = True
    
    caracteristicas = [
        ('Enfoque', 'Ágil'),
        ('Modelo', 'Incremental'),
        ('Metodología', 'Extreme Programming (XP)')
    ]
    
    for caract, valor in caracteristicas:
        p = doc.add_paragraph()
        run1 = p.add_run(f'{caract}: ')
        run1.bold = True
        run2 = p.add_run(valor)
    
    doc.add_paragraph()
    
    # Dimensiones del modelo Boehm-Turner
    dimensiones = [
        ('Criticidad — Baja (riesgo bajo)', 'El proyecto no representa un riesgo significativo para las operaciones del cliente ni para la continuidad del servicio que ofrece. Un posible fallo del sistema no generaría pérdidas graves ni comprometería la reputación del negocio. Esta condición permite manejar el desarrollo con una planificación flexible, priorizando la eficiencia en el uso de recursos sin comprometer la estabilidad operativa.'),
        ('Dinamismo — Medio (requisitos con cambios moderados)', 'El contexto comercial del cliente puede variar ocasionalmente, generando ajustes en los requerimientos del sistema. Estos cambios no son constantes ni impredecibles, pero deben considerarse dentro de la gestión del proyecto. A nivel de negocio, esto significa que el equipo debe mantener capacidad de respuesta ante modificaciones razonables, equilibrando la estabilidad del proyecto con la necesidad de adaptarse a nuevas demandas.'),
        ('Cultura — Baja o nula (equipo con poca colaboración interna)', 'El equipo de desarrollo no posee una cultura consolidada de trabajo colaborativo. La comunicación interna es limitada y las decisiones tienden a tomarse de forma individual. A nivel de negocio, esto puede impactar la coordinación, el seguimiento de tareas y la velocidad de respuesta ante imprevistos. Es necesario reforzar la organización interna y establecer mecanismos claros de comunicación para mejorar la eficiencia global del proyecto.'),
        ('Tamaño del equipo — Pequeño (2 personas)', 'El equipo está compuesto por dos integrantes, lo que permite un control directo de las actividades y una comunicación sencilla. Sin embargo, este tamaño también implica una menor capacidad para distribuir las cargas de trabajo o cubrir ausencias. Desde la perspectiva del negocio, se requiere una planificación precisa y una clara asignación de responsabilidades para evitar retrasos y mantener la continuidad del proyecto.'),
        ('Personal — Dos desarrolladores junior (experiencia baja)', 'El equipo está integrado por personal en formación, con poca experiencia práctica en proyectos reales. Esto representa un factor relevante en la gestión del negocio, ya que puede influir en los tiempos de entrega y en la calidad de los resultados. Es necesario prever instancias de revisión, apoyo y orientación para asegurar que el trabajo cumpla los objetivos planteados y mantenga un estándar de calidad adecuado.')
    ]
    
    for titulo, descripcion in dimensiones:
        p = doc.add_paragraph()
        run_titulo = p.add_run(titulo)
        run_titulo.bold = True
        doc.add_paragraph(descripcion)
    
    # 4.1.2 Modelo de proceso de software
    doc.add_paragraph()
    p = doc.add_paragraph('4.1.2 Modelo de proceso de software')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_1_2 = """Se adopta un modelo de proceso incremental, implementado mediante la metodología Extreme Programming (XP). En este enfoque, el desarrollo se organiza en iteraciones cortas (de una a dos semanas) que permiten construir y agregar al sistema nuevos incrementos funcionales. Cada incremento representa una parte operativa del producto, completamente integrada y validada."""
    
    doc.add_paragraph(texto_4_1_2)
    
    doc.add_paragraph()
    texto_4_1_2_cont = """Esta estrategia combina la planificación continua con la entrega temprana de valor, lo que facilita validar las funcionalidades implementadas con el cliente y adaptarse con agilidad a posibles cambios en los requisitos. XP proporciona las prácticas técnicas que sustentan este modelo incremental: planificación basada en historias de usuario, programación en parejas, desarrollo guiado por pruebas (TDD), integración continua y refactorización constante. Estas prácticas garantizan que cada incremento mantenga un nivel adecuado de calidad y que el sistema evolucione de forma controlada, verificable y alineada con los objetivos del proyecto."""
    
    doc.add_paragraph(texto_4_1_2_cont)
    
    # 4.1.3 Método de ingeniería de software
    doc.add_paragraph()
    p = doc.add_paragraph('4.1.3 Método de ingeniería de software (Extreme Programming)')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_1_3 = """Se selecciona Extreme Programming (XP) como método principal del proceso de desarrollo, dada su orientación hacia equipos pequeños, entregas frecuentes y adaptación rápida a los cambios. La metodología XP integra valores fundamentales como la comunicación, simplicidad, retroalimentación y adaptabilidad del equipo, los cuales se ajustan al entorno dinámico de una MIPYME que requiere optimizar procesos comerciales sin grandes estructuras tecnológicas."""
    
    doc.add_paragraph(texto_4_1_3)
    
    doc.add_paragraph()
    p = doc.add_paragraph('Las prácticas principales aplicadas fueron:')
    p.runs[0].bold = True
    
    practicas_xp = [
        'Desarrollo iterativo con ciclos cortos de liberación (1 a 2 semanas).',
        'Retroalimentación continua del cliente para validar funcionalidades de pedidos, pagos y clientes.',
        'Programación en parejas para mejorar la calidad del código y reducir errores.',
        'Integración continua, garantizando el funcionamiento correcto del sistema tras cada iteración.',
        'Refactorización constante, mejorando la estructura interna sin alterar la funcionalidad.',
        'Pruebas automatizadas para asegurar la fiabilidad de las operaciones críticas.'
    ]
    
    for practica in practicas_xp:
        doc.add_paragraph(practica, style='List Bullet')
    
    doc.add_paragraph()
    texto_complemento = """El marco metodológico de XP se complementó con herramientas y prácticas específicas de ingeniería para mantener la calidad del producto y la trazabilidad de los procesos: planificación de iteraciones con historias de usuario priorizadas según el valor de negocio, diseño orientado a componentes diferenciando claramente los módulos de clientes, pedidos, productos y pagos, desarrollo incremental liberando funcionalidades operativas al cierre de cada iteración, revisiones técnicas continuas y pruebas unitarias en cada módulo, e integración y despliegue continuo mediante un entorno de control de versiones."""
    
    doc.add_paragraph(texto_complemento)
    
    doc.add_paragraph()
    texto_artefactos = """Los artefactos principales generados durante el proceso incluyen: historias de usuario, versión funcional del sistema en cada iteración, documentación técnica y registros de pruebas."""
    
    doc.add_paragraph(texto_artefactos)
    
    # 4.2 Herramientas y tecnologías
    doc.add_paragraph()
    p = doc.add_paragraph('4.2 Herramientas y tecnologías')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    # 4.2.1 Herramienta CASE
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.1 Herramienta CASE')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_1 = """Visual Paradigm 8.0 y Lucidchart se emplearon para el modelado de diagramas UML y la representación de la arquitectura del sistema. Estas herramientas permiten la colaboración en tiempo real, la documentación visual del flujo de procesos y la integración directa con entornos de desarrollo, lo que facilitó la comunicación y coordinación dentro del equipo de trabajo."""
    
    doc.add_paragraph(texto_4_2_1)
    
    # 4.2.2 Lenguaje de modelado
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.2 Lenguaje de modelado')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_2 = """Se utilizó UML (Unified Modeling Language) en su versión 2.0 como estándar para la representación gráfica de sistemas de software. Esta herramienta permitió describir de forma estructurada los casos de uso, clases, secuencias y componentes, facilitando la comprensión del sistema tanto a nivel funcional como técnico."""
    
    doc.add_paragraph(texto_4_2_2)
    
    # 4.2.3 Marco de trabajo para el desarrollo
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.3 Marco de trabajo para el desarrollo')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_3 = """El desarrollo de la aplicación se realizó utilizando el framework FastAPI para el backend, complementado con HTML5, CSS3 y JavaScript ES2023 para el frontend. FastAPI fue seleccionado por su alto rendimiento, generación automática de documentación interactiva mediante OpenAPI, validación de datos con Pydantic, soporte nativo para operaciones asíncronas y facilidad de integración con ORMs como SQLAlchemy. En conjunto, estas tecnologías permitieron construir una plataforma web dinámica, segura y modular, adecuada para la gestión de pedidos, clientes, productos y pagos en la MIPYME."""
    
    doc.add_paragraph(texto_4_2_3)
    
    # 4.2.4 Entorno de desarrollo integrado
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.4 Entorno de desarrollo integrado')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_4 = """Se empleó Visual Studio Code como entorno de desarrollo principal, por su compatibilidad con Python y JavaScript, integración con Git, extensiones para depuración, control de versiones y pruebas, y soporte para desarrollo remoto mediante SSH. Este entorno favoreció la productividad del equipo y la trazabilidad de los avances en las iteraciones de XP."""
    
    doc.add_paragraph(texto_4_2_4)
    
    # 4.2.5 Lenguaje de programación
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.5 Lenguaje de programación')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_5 = """El sistema fue implementado utilizando Python 3.12 como lenguaje principal, aprovechando su sintaxis legible, amplia comunidad, ecosistema de librerías orientadas al desarrollo web y soporte para programación asíncrona. En el lado del cliente, se utilizó JavaScript ES2023 para la creación de interfaces interactivas y la validación dinámica de formularios, garantizando una experiencia de usuario fluida y moderna."""
    
    doc.add_paragraph(texto_4_2_5)
    
    # 4.2.6 Gestor de base de datos
    doc.add_paragraph()
    p = doc.add_paragraph('4.2.6 Gestor de base de datos')
    p.runs[0].bold = True
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(12)
    
    texto_4_2_6 = """Se seleccionó PostgreSQL 16 como sistema gestor de base de datos por su cumplimiento de las propiedades ACID, soporte para tipos de datos JSON, funciones definidas por el usuario (como calcular_monto_pendiente), triggers para automatización de reglas de negocio, alto rendimiento y robustez en entornos transaccionales. El control y la administración del modelo de datos se realizaron mediante pgAdmin 4, que facilitó la creación, consulta y mantenimiento de las estructuras de datos que sustentan la aplicación."""
    
    doc.add_paragraph(texto_4_2_6)
    
    # Conclusiones parciales del Epígrafe IV
    doc.add_paragraph()
    p = doc.add_paragraph('Conclusiones parciales')
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(12)
    
    conclusiones_epi4 = """La aplicación del modelo de Boehm y Turner justificó la adopción de la metodología Extreme Programming (XP) como el enfoque más adecuado para un equipo pequeño (2 personas), con criticidad baja, dinamismo medio, cultura limitada de colaboración y personal junior, posibilitando un desarrollo incremental, iterativo y flexible. La selección de un modelo de proceso incremental con iteraciones cortas (1-2 semanas) permitió entregar valor temprano al cliente y adaptarse ágilmente a cambios en los requisitos. La adopción de XP como método de ingeniería de software garantizó la aplicación de prácticas técnicas rigurosas: desarrollo iterativo, retroalimentación continua, programación en parejas, integración continua, refactorización constante y pruebas automatizadas. Asimismo, la selección de herramientas y tecnologías conformó un entorno moderno, seguro y escalable, sustentado en FastAPI, Python 3.12, PostgreSQL 16, JavaScript ES2023, Visual Studio Code, UML 2.0 y Visual Paradigm, garantizando la integridad de los datos, la mantenibilidad del sistema y su adaptabilidad futura. En conjunto, este epígrafe establece los fundamentos metodológicos y tecnológicos que sirvieron de base para el diseño y la implementación del sistema, en coherencia con las mejores prácticas de la ingeniería de software aplicada a entornos empresariales."""
    
    doc.add_paragraph(conclusiones_epi4)
    
    # Salto de página
    doc.add_page_break()

def main():
    """Función principal que genera el documento"""
    import os
    filename = 'Informe_Tecnico_PID_NUEVO_V1.docx'
    
    # Verificar si el documento ya existe
    if os.path.exists(filename):
        print(f"📄 Abriendo documento existente: {filename}")
        print("⚠️  SOLO se agregará el nuevo contenido al final, preservando tus cambios manuales.")
        doc = Document(filename)
        
        # Agregar solo el nuevo epígrafe (comentar los que ya fueron agregados)
        # print("Agregando Epígrafe V...")
        # epigrafe_v(doc)
        
        print("\n✅ Documento actualizado exitosamente.")
        print("📝 Para agregar un nuevo epígrafe, descomenta la línea correspondiente en main().")
        
    else:
        print("Generando documento nuevo con primeras secciones...")
        
        # Crear documento
        doc = Document()
        
        # Configurar márgenes y estilos
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1)
        
        # Crear estilos personalizados
        crear_estilos(doc)
        
        # Generar secciones
        print("  1. Portada...")
        portada(doc)
        
        print("  2. Resumen...")
        resumen(doc)
        
        print("  3. Abstract...")
        abstract(doc)
        
        print("  4. Placeholders de índices...")
        indices_placeholders(doc)
        
        print("  5. Placeholder de opinión del tutor...")
        opinion_tutor_placeholder(doc)
        
        print("  6. Introducción...")
        introduccion(doc)
        
        print("  7. DESARROLLO (intro + Epígrafe I)...")
        desarrollo(doc)
        
        print("  8. DESARROLLO (Epígrafe II)...")
        epigrafe_ii(doc)
        
        print("  9. DESARROLLO (Epígrafe III)...")
        epigrafe_iii(doc)
        
        print(" 10. DESARROLLO (Epígrafe IV)...")
        epigrafe_iv(doc)
        
        print(f"\n✅ Documento generado exitosamente: {filename}")
        print(f"📄 Revisa el documento y aprueba antes de continuar con los capítulos.")
    
    # Guardar documento
    doc.save(filename)
    print(f"💾 Guardado: {filename}")
    print(f"📚 Consulta INSTRUCCIONES_PROYECTO_COMPLETO.md para toda la información del proyecto.")

if __name__ == "__main__":
    main()
